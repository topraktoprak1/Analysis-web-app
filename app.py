from flask import Flask, render_template, request, jsonify, send_file, session, redirect, url_for
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import json
import os
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from functools import wraps

app = Flask(__name__)

# Security configuration
import secrets
app.secret_key = os.environ.get('SECRET_KEY', secrets.token_hex(32))
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Session security
app.config['SESSION_COOKIE_SECURE'] = False  # Set True in production with HTTPS
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['PERMANENT_SESSION_LIFETIME'] = 3600  # 1 hour

# PostgreSQL configuration
# For local development, use: postgresql://username:password@localhost:5432/database_name
# For production, use environment variable
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 
    'postgresql://postgres:1234@localhost:8080/veri_analizi')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'pool_size': 10,
    'pool_recycle': 3600,
    'pool_pre_ping': True,
}

db = SQLAlchemy(app)

# Cache for data to avoid reloading from database every time
_data_cache = {
    'data': None,
    'timestamp': None,
    'record_count': 0
}

def get_cache_key():
    """Get cache key based on database record count"""
    count = DatabaseRecord.query.count()
    return count

def clear_data_cache():
    """Clear the data cache"""
    global _data_cache
    _data_cache = {'data': None, 'timestamp': None, 'record_count': 0}

# Database Models
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)
    name = db.Column(db.String(120), nullable=False)
    role = db.Column(db.String(20), default='user')
    email = db.Column(db.String(120))
    first_name = db.Column(db.String(80))
    last_name = db.Column(db.String(80))
    profile_photo = db.Column(db.String(200), default='img/avatars/avatar1.jpeg')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class DatabaseRecord(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    personel = db.Column(db.String(120), nullable=False)
    data = db.Column(db.Text, nullable=False)  # JSON string of record data
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

class SavedFilter(db.Model):
    """Model to store user's saved filter configurations"""
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    filter_name = db.Column(db.String(200), nullable=False)
    filter_type = db.Column(db.String(50), nullable=False)  # 'database', 'pivot', or 'graph'
    filter_config = db.Column(db.Text, nullable=False)  # JSON string of filter configuration
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relationship
    user = db.relationship('User', backref='saved_filters')

# Initialize database and create admin user
def init_db():
    db.create_all()
    # Create admin if not exists
    admin = User.query.filter_by(username='admin').first()
    if not admin:
        admin = User(
            username='admin',
            password=generate_password_hash('admin123'),
            name='Administrator',
            role='admin'
        )
        db.session.add(admin)
        db.session.commit()
        print('✓ Admin user created: admin/admin123')
    else:
        print('✓ Database initialized')

# Login required decorator
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            return redirect(url_for('login_page'))
        return f(*args, **kwargs)
    return decorated_function

# Admin required decorator
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session or session.get('role') != 'admin':
            return jsonify({'error': 'Admin access required'}), 403
        return f(*args, **kwargs)
    return decorated_function

# Create uploads folder if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Favorites file path
FAVORITES_FILE_PATH = os.path.join(os.path.dirname(__file__), 'favorite_reports.json')

# Utility functions
def load_excel_data(file_path, user_filter=None):
    """Load Excel file and return DataFrame from DATABASE sheet with optional user filtering"""
    file_name = os.path.basename(file_path).lower()
    
    if file_name.endswith('.xlsb'):
        try:
            df = pd.read_excel(file_path, sheet_name='DATABASE', engine='pyxlsb')
       
        except:
            df = pd.read_excel(file_path, engine='pyxlsb')
    else:
        try:
            df = pd.read_excel(file_path, sheet_name='DATABASE')
        except:
            df = pd.read_excel(file_path)
    
    print(f"Loaded Excel with {len(df)} rows and {len(df.columns)} columns")
    
    # Clean column names - remove newlines and extra spaces
    df.columns = df.columns.str.replace('\n', ' ').str.replace('\r', ' ').str.replace('  ', ' ').str.strip()
    print(f"Cleaned column names")
    
    # Preserve date formats - convert datetime columns to string in dd/mmm/yyyy format
    for col in df.columns:
        # Check if column name contains date-related keywords (case insensitive, including parentheses)
        col_lower = str(col).lower() if col else ''
        if any(keyword in col_lower for keyword in ['week', 'month', 'date', 'tarih']):
            try:
                print(f"\nProcessing date column: '{col}'")
                print(f"Column dtype: {df[col].dtype}")
                print(f"First 5 raw values: {df[col].head(5).tolist()}")
                import sys
                sys.stdout.flush()  # Force output to appear immediately
                
                # If already datetime, just format it
                if df[col].dtype == 'datetime64[ns]':
                    # Manual formatting to ensure year is included
                    df[col] = df[col].apply(lambda x: x.strftime('%d/%b/%Y') if pd.notna(x) else '')
                    print(f"Converted datetime column directly")
                    print(f"First 5 after conversion: {df[col].head(5).tolist()}")
                else:
                    # Check if values are numbers (Excel date serial numbers)
                    first_val = df[col].dropna().iloc[0] if len(df[col].dropna()) > 0 else None
                    
                    if first_val is not None and isinstance(first_val, (int, float)) and first_val > 30000:
                        # These are Excel serial date numbers, convert them
                        print(f"Detected Excel serial numbers, converting...")
                        # Excel dates start from 1900-01-01 (serial 1)
                        # Convert Excel serial dates to pandas datetime
                        from datetime import datetime, timedelta
                        excel_start = datetime(1899, 12, 30)
                        date_series = df[col].apply(lambda x: excel_start + timedelta(days=float(x)) if pd.notna(x) and isinstance(x, (int, float)) else pd.NaT)
                        print(f"Successfully parsed {date_series.notna().sum()} out of {len(date_series)} values")
                        
                        # Format the dates manually to ensure year is included
                        def format_date(dt):
                            if pd.notna(dt):
                                try:
                                    # Manual string formatting
                                    day = dt.day
                                    month = dt.strftime('%b')
                                    year = dt.year
                                    formatted = f"{day:02d}/{month}/{year}"
                                    return formatted
                                except:
                                    return ''
                            return ''
                        
                        df[col] = date_series.apply(format_date)
                        print(f"First 5 after formatting: {df[col].head(5).tolist()}")
                        
                        # Extra debug: Check if year is actually there
                        sample_date = date_series.dropna().iloc[0] if len(date_series.dropna()) > 0 else None
                        if sample_date:
                            print(f"DEBUG: Sample datetime object - Day:{sample_date.day}, Month:{sample_date.month}, Year:{sample_date.year}")
                            print(f"DEBUG: Formatted sample: {format_date(sample_date)}")
                    else:
                        # Try standard datetime parsing
                        original_values = df[col].astype(str).copy()
                        date_series = pd.to_datetime(df[col], errors='coerce', format='mixed')
                        
                        success_count = date_series.notna().sum()
                        print(f"Successfully parsed {success_count} out of {len(date_series)} values")
                        
                        if success_count > 0:
                            formatted = date_series.dt.strftime('%d/%b/%Y')
                            df[col] = formatted.where(date_series.notna(), original_values)
                            print(f"First 5 after formatting: {df[col].head(5).tolist()}")
                        else:
                            print(f"No dates could be parsed, keeping original values")
                        
            except Exception as e:
                print(f"Error processing date column {col}: {e}")
                import traceback
                traceback.print_exc()
    
    # Filter by user if not admin
    if user_filter and 'PERSONEL' in df.columns:
        df = df[df['PERSONEL'].str.strip().str.upper() == user_filter.strip().upper()]
    
    return df

def get_data_from_db(user_filter=None):
    """Get data from database records with caching"""
    global _data_cache
    
    try:
        # DISABLE CACHE TEMPORARILY FOR DEBUGGING
        # Check if we can use cached data (only for no filter case)
        current_count = DatabaseRecord.query.count()
        
        print(f"DEBUG: Database has {current_count} records")
        
        # FORCE RELOAD - skip cache
        # if user_filter is None and _data_cache['data'] is not None and _data_cache['record_count'] == current_count:
        #     print(f"Using cached data ({current_count} records, {len(_data_cache['data'])} rows)")
        #     return _data_cache['data'].copy()
        
        # Load fresh data from database
        print(f"Loading data from database ({current_count} records)...")
        if user_filter:
            records = DatabaseRecord.query.filter_by(personel=user_filter).all()
            print(f"DEBUG: Filtered to {len(records)} records for user {user_filter}")
        else:
            records = DatabaseRecord.query.all()
            print(f"DEBUG: Loaded all {len(records)} records")
        
        if not records:
            print("DEBUG: No records found")
            return pd.DataFrame()
        
        records_list = []
        for record in records:
            try:
                record_dict = json.loads(record.data)
                # Debug: Check date format after loading from database
                if len(records_list) == 0 and '(Week / Month)' in record_dict:
                    print(f"DEBUG LOAD: First date value after DB load: {record_dict['(Week / Month)']}")
                records_list.append(record_dict)
            except:
                continue
        
        print(f"DEBUG: Parsed {len(records_list)} records into DataFrame")
        
        if records_list:
            # Create DataFrame without automatic date parsing
            df = pd.DataFrame(records_list, dtype=object)
            
            # Ensure date columns stay as strings (don't let pandas auto-convert)
            for col in df.columns:
                if 'week' in str(col).lower() or 'month' in str(col).lower() or 'date' in str(col).lower():
                    df[col] = df[col].astype(str)
            
            # DEBUG: Force print to console
            import sys
            print("\n" + "="*80, file=sys.stderr, flush=True)
            print(f"DEBUG DF: Created DataFrame with {len(df)} rows and {len(df.columns)} columns", file=sys.stderr, flush=True)
            
            # Debug: Check date column in DataFrame
            if '(Week / Month)' in df.columns:
                first_dates = df['(Week / Month)'].head(5).tolist()
                print(f"DEBUG DF: (Week / Month) first 5 values: {first_dates}", file=sys.stderr, flush=True)
                print(f"DEBUG DF: Sample date length: {len(first_dates[0]) if first_dates else 0}", file=sys.stderr, flush=True)
            print("="*80 + "\n", file=sys.stderr, flush=True)
            
            # Ensure PERSONEL column exists for filtering but don't expose it
            if 'Name Surname' in df.columns and 'PERSONEL' not in df.columns:
                df['PERSONEL'] = df['Name Surname']
            elif 'PERSONEL' in df.columns and 'Name Surname' not in df.columns:
                # If only PERSONEL exists, create Name Surname
                df['Name Surname'] = df['PERSONEL']
            
            # Cache the data if no filter
            if user_filter is None:
                _data_cache['data'] = df.copy()
                _data_cache['timestamp'] = datetime.now()
                _data_cache['record_count'] = current_count
                print(f"Data cached successfully")
            
            return df
        return pd.DataFrame()
    except Exception as e:
        print(f"Database load error: {str(e)}")
        return pd.DataFrame()

def get_combined_data(file_path=None, user_filter=None):
    """Get data from both Excel file and database"""
    dfs = []
    
    # Load from Excel if file exists
    if file_path and os.path.exists(file_path):
        try:
            df_excel = load_excel_data(file_path, user_filter)
            if not df_excel.empty:
                dfs.append(df_excel)
        except:
            pass
    
    # Load from database
    df_db = get_data_from_db(user_filter)
    if not df_db.empty:
        dfs.append(df_db)
    
    # Combine dataframes
    if dfs:
        combined_df = pd.concat(dfs, ignore_index=True)
        return combined_df
    return pd.DataFrame()

# Global cache for lookup tables
_lookup_cache = {'info': None, 'rates': None, 'timestamp': None}

import os
import pandas as pd

# initialize cache at module level
_lookup_cache = {'info': None, 'rates': None, 'timestamp': None}

def _find_column(df, *must_have_substrings):
    """Return first column name in df whose lowercase name contains all substrings"""
    for col in df.columns:
        name = str(col).lower()
        if all(s.lower() in name for s in must_have_substrings):
            return col
    return None

def load_lookup_tables(file_path=None):
    """Load Info and Hourly Rates sheets for XLOOKUP operations (robust, cached)."""
    global _lookup_cache

    try:
        # if file_path not provided, pick newest .xlsb in UPLOAD_FOLDER
        if not file_path:
            upload_dir = app.config.get('UPLOAD_FOLDER')  # safer access
            if not upload_dir or not os.path.isdir(upload_dir):
                print("Warning: UPLOAD_FOLDER not set or invalid")
                return None, None
            xlsb_files = [f for f in os.listdir(upload_dir) if f.lower().endswith('.xlsb')]
            if not xlsb_files:
                return None, None
            xlsb_files.sort(reverse=True)
            file_path = os.path.join(upload_dir, xlsb_files[0])

        if not os.path.isfile(file_path):
            print(f"Warning: lookup file not found: {file_path}")
            return None, None

        file_mtime = os.path.getmtime(file_path)

        # validate cache
        if (_lookup_cache.get('timestamp') == file_mtime and
                _lookup_cache.get('info') is not None and
                _lookup_cache.get('rates') is not None):
            return _lookup_cache['info'], _lookup_cache['rates']

        # read sheets - allow header variations; keep as DataFrame even if small
        df_info = pd.read_excel(file_path, sheet_name='Info', engine='pyxlsb')
        # 'Hourly Rates' often has a header row offset; try header=0 then header=1 fallback
        try:
            df_rates = pd.read_excel(file_path, sheet_name='Hourly Rates', engine='pyxlsb', header=0)
        except Exception:
            df_rates = pd.read_excel(file_path, sheet_name='Hourly Rates', engine='pyxlsb', header=1)

        _lookup_cache = {'info': df_info, 'rates': df_rates, 'timestamp': file_mtime}
        print(f"Loaded lookup tables: Info ({len(df_info)} rows), Rates ({len(df_rates)} rows)")
        return df_info, df_rates

    except Exception as e:
        print(f"Warning: Could not load lookup tables: {e}")
        return None, None


def xlookup(lookup_value, lookup_array, return_array, default=None):
    """
    Simulate Excel XLOOKUP:
     - works with pandas Series / lists / numpy arrays
     - tries numeric match first (when both sides numeric), then case-insensitive string exact match
     - returns 'default' when not found or on error
    """
    try:
        # normalize arrays to pandas Series
        if not isinstance(lookup_array, pd.Series):
            lookup_array = pd.Series(list(lookup_array))
        if not isinstance(return_array, pd.Series):
            return_array = pd.Series(list(return_array))

        # empty
        if lookup_array.empty:
            return default

        # empty lookup value
        if lookup_value is None or (isinstance(lookup_value, str) and lookup_value.strip() == '') or pd.isna(lookup_value):
            return default

        # attempt numeric comparison: convert both to numeric where possible
        lookup_num = pd.to_numeric(pd.Series([lookup_value]), errors='coerce').iloc[0]
        arr_nums = pd.to_numeric(lookup_array, errors='coerce')

        if not pd.isna(lookup_num):
            # find exact numeric match (positional)
            matches = arr_nums[arr_nums.notna() & (arr_nums == lookup_num)]
            if not matches.empty:
                pos = matches.index[0]  # label of first match
                # use .loc to preserve label alignment between lookup_array and return_array
                if pos in return_array.index:
                    result = return_array.loc[pos]
                else:
                    # fallback to positional iloc
                    result = return_array.iloc[pos] if pos < len(return_array) else default
                return default if pd.isna(result) else result

        # fallback to case-insensitive string match
        lookup_str = str(lookup_value).strip().upper()
        # create upper-version of lookup_array values (skip NaN)
        for pos, val in lookup_array.items():
            if pd.isna(val):
                continue
            if str(val).strip().upper() == lookup_str:
                result = return_array.loc[pos] if pos in return_array.index else return_array.iloc[pos]
                return default if pd.isna(result) else result

        return default

    except Exception as e:
        print(f"XLOOKUP error for value '{lookup_value}': {e}")
        return default


def add_calculated_columns(df):
    """Add calculated columns based on Excel formulas (robust variant)."""
    if df is None:
        return df
    if df.empty:
        return df.copy()

    df = df.copy()

    # load lookup tables
    df_info, df_rates = load_lookup_tables()

    if df_info is not None:
        print(f"Info sheet columns: {list(df_info.columns)[:60]}")
    if df_rates is not None:
        print(f"Rates sheet columns: {list(df_rates.columns)[:20]}")

    # === 23. LS/Unit Rate - FIRST used by other formulas ===
    def calc_ls_unit(row):
        proj_group = str(row.get('Projects/Group', '')).lower()
        company = str(row.get('Company', '')).lower()
        # if "lumpsum" appears anywhere in project/group name -> Lumpsum
        if 'lumpsum' in proj_group:
            return 'Lumpsum'
        # company checks (normalize)
        if any(k in company for k in ['i4', 'degenkolb', 'kilci danışmanlık', 'kilci danismanlik', 'kilci']):
            return 'Lumpsum'
        return 'Unit Rate'

    if 'Projects/Group' in df.columns or 'Company' in df.columns:
        df['LS/Unit Rate'] = df.apply(calc_ls_unit, axis=1)

    # === 7. AP-CB/Subcon ===
    if 'Company' in df.columns:
        df['AP-CB/Subcon'] = df['Company'].apply(
            lambda x: 'AP-CB' if 'AP-CB' in str(x).upper() else 'Subcon'
        )

    # === 1. North/South from Info sheet ===
    if df_info is not None and 'Projects/Group' in df.columns:
        # try to find a column in Info that indicates North/South
        north_south_col = None
        for col in df_info.columns:
            if 'north' in str(col).lower() or 'south' in str(col).lower():
                north_south_col = col
                break
        # find a Projects/Group-like column in Info
        proj_group_col = _find_column(df_info, 'project', 'group') or df_info.columns[0]
        if north_south_col:
            df['North/South'] = df['Projects/Group'].apply(
                lambda x: xlookup(x, df_info[proj_group_col], df_info[north_south_col], '')
            )

    # === 2. Currency ===
    if df_rates is not None:
        emp_col = _find_column(df_rates, 'emp') or _find_column(df_rates, 'name', 'surname')
        currency_col = _find_column(df_rates, 'currency') or _find_column(df_rates, 'curr')

        if 'Emp No' in df.columns and emp_col and currency_col:
            df['Currency'] = df['Emp No'].apply(
                lambda x: 'TL' if str(x).strip() == '905264' else xlookup(x, df_rates[emp_col], df_rates[currency_col], 'USD')
            )
        elif 'Name Surname' in df.columns and emp_col and currency_col:
            df['Currency'] = df['Name Surname'].apply(
                lambda x: xlookup(x, df_rates[emp_col], df_rates[currency_col], 'USD')
            )

    # === 5. Hourly Base Rate ===
    if df_rates is not None:
        emp_col = _find_column(df_rates, 'emp') or _find_column(df_rates, 'name', 'surname')
        # try to find base rate columns (may be named 'Hourly Base Rate', 'Base Rate', etc.)
        base_rate_cols = [col for col in df_rates.columns if ('hourly' in str(col).lower() and 'base' in str(col).lower()) or ('base' in str(col).lower() and 'rate' in str(col).lower())]
        # fallback to any numeric-looking rate columns
        if not base_rate_cols:
            base_rate_cols = [col for col in df_rates.columns if 'rate' in str(col).lower()]

        if emp_col and base_rate_cols:
            def calc_base_rate(row):
                lookup_val = row.get('Emp No') or row.get('Name Surname')
                if not lookup_val:
                    return 0.0
                ap_cb_subcon = row.get('AP-CB/Subcon', '')
                ls_unit = row.get('LS/Unit Rate', '')
                # choose a second candidate if available for Subcon+Unit Rate condition
                if len(base_rate_cols) > 1 and ap_cb_subcon == 'Subcon' and ls_unit == 'Unit Rate':
                    rate_col = base_rate_cols[1]
                else:
                    rate_col = base_rate_cols[0]
                result = xlookup(lookup_val, df_rates[emp_col], df_rates[rate_col], 0)
                try:
                    return float(result) if result not in (None, '') else 0.0
                except Exception:
                    return 0.0

            df['Hourly Base Rate'] = df.apply(calc_base_rate, axis=1)

    # === 6. Hourly Additional Rate ===
    if df_rates is not None:
        emp_col = _find_column(df_rates, 'emp') or _find_column(df_rates, 'name', 'surname')
        additional_rate_col = _find_column(df_rates, 'additional', 'rate') or _find_column(df_rates, 'extra', 'rate')
        if emp_col and additional_rate_col:
            def calc_additional_rate(row):
                if row.get('LS/Unit Rate') == 'Lumpsum':
                    return 0.0
                if 'AP-CB' in str(row.get('Company', '')).upper():
                    return 0.0
                lookup_val = row.get('Emp No') or row.get('Name Surname')
                if not lookup_val:
                    return 0.0
                additional = xlookup(lookup_val, df_rates[emp_col], df_rates[additional_rate_col], 0) or 0
                try:
                    additional = float(additional)
                except Exception:
                    additional = 0.0

                currency = row.get('Currency', '')
                if currency == 'USD' or currency == '' or currency is None:
                    return additional
                elif currency == 'TL' and df_info is not None:
                    # try to find week/month and TL rate columns dynamically in Info
                    week_col = _find_column(df_info, 'week') or _find_column(df_info, 'month') or df_info.columns[0]
                    tl_rate_col = None
                    for c in df_info.columns:
                        cl = str(c).lower()
                        if ('tl' in cl or 'usd' in cl) and ('rate' in cl or 'kur' in cl):
                            # prefer TL-specific
                            if 'tl' in cl:
                                tl_rate_col = c
                                break
                            tl_rate_col = c
                    if week_col and tl_rate_col:
                        week_month = row.get('(Week / Month)') or row.get('Week / Month') or ''
                        exchange_rate = xlookup(week_month, df_info[week_col], df_info[tl_rate_col], 1) or 1
                        try:
                            # Convert TL additional to USD by dividing by TL→USD rate
                            rate = float(exchange_rate) if exchange_rate else 1.0
                            return additional / rate if rate else additional
                        except Exception:
                            return additional
                else:
                    return additional

            df['Hourly Additional Rate'] = df.apply(calc_additional_rate, axis=1)

    # === 3. Hourly Rate ===
    if 'Hourly Base Rate' in df.columns or 'Hourly Additional Rate' in df.columns:
        df['Hourly Base Rate'] = pd.to_numeric(df.get('Hourly Base Rate', 0), errors='coerce').fillna(0)
        df['Hourly Additional Rate'] = pd.to_numeric(df.get('Hourly Additional Rate', 0), errors='coerce').fillna(0)
        df['Hourly Rate'] = df['Hourly Base Rate'] + df['Hourly Additional Rate']

    # === Helper: normalize week/month label for lookups ===
    def _normalize_week_label(val):
        try:
            if pd.isna(val):
                return ''
            s = str(val).strip()
            # If looks like a date, convert to dd/Mon/YYYY
            dt = pd.to_datetime(s, errors='coerce')
            if pd.notna(dt):
                return dt.strftime('%d/%b/%Y')
            # If looks like "W46" or "Week 46" normalize to "W46"
            import re
            m = re.search(r'(?:w|week)\s*(\d{1,2})', s, re.IGNORECASE)
            if m:
                return f"W{int(m.group(1))}"
            return s
        except Exception:
            return str(val) if val is not None else ''

    # === 8. General Total Cost (USD) ===
    if 'Currency' in df.columns and 'Cost' in df.columns and df_info is not None:
        # try to locate week/month column and exchange columns in Info
        week_col = _find_column(df_info, 'week') or _find_column(df_info, 'month') or (df_info.columns[0] if len(df_info.columns) else None)
        usd_rate_col = _find_column(df_info, 'usd', 'rate') or _find_column(df_info, 'usd')
        euro_rate_col = _find_column(df_info, 'euro', 'rate') or _find_column(df_info, 'euro')

        def calc_general_cost(row):
            currency = str(row.get('Currency', '')).upper()
            cost = pd.to_numeric(row.get('Cost', 0), errors='coerce') or 0
            week_month = _normalize_week_label(row.get('(Week / Month)') or row.get('Week / Month') or '')
            try:
                if currency == 'TL':
                    if week_col and usd_rate_col:
                        exchange_rate = xlookup(week_month, df_info[week_col], df_info[usd_rate_col], 1) or 1
                        return cost / float(exchange_rate) if exchange_rate else cost
                    return cost
                elif currency in ('EURO', 'EUR'):
                    if week_col and euro_rate_col:
                        exchange_rate = xlookup(week_month, df_info[week_col], df_info[euro_rate_col], 1) or 1
                        return cost * float(exchange_rate) if exchange_rate else cost
                    return cost
                else:
                    return cost
            except Exception:
                return cost

        df['General Total Cost (USD)'] = df.apply(calc_general_cost, axis=1)

    # === 9. Hourly Unit Rate (USD) ===
    if 'General Total Cost (USD)' in df.columns and 'Total Hours' in df.columns:
        # avoid division by zero by producing NaN if Total Hours zero/NA
        df['Hourly Unit Rate (USD)'] = pd.to_numeric(df['General Total Cost (USD)'], errors='coerce') / pd.to_numeric(df['Total Hours'], errors='coerce')
        df['Hourly Unit Rate (USD)'] = df['Hourly Unit Rate (USD)'].replace([pd.NA, pd.NaT], float('nan'))

    # === 10-13 (simplified fallback versions) ===
    # Implementing simpler versions since Summary sheet structure unknown.
    if 'Hourly Unit Rate (USD)' in df.columns:
        df['İşveren Hakediş Birim Fiyat (USD)'] = df['Hourly Unit Rate (USD)']

    if 'İşveren Hakediş Birim Fiyat (USD)' in df.columns:
        if 'Special Hours' in df.columns:
            df['İşveren-Hakediş(USD)'] = df.apply(
                lambda row: (pd.to_numeric(row.get('Special Hours', 0), errors='coerce') * pd.to_numeric(row.get('İşveren Hakediş Birim Fiyat (USD)', 0), errors='coerce'))
                if pd.to_numeric(row.get('Special Hours', 0), errors='coerce') > 0
                else (pd.to_numeric(row.get('İşveren Hakediş Birim Fiyat (USD)', 0), errors='coerce') * pd.to_numeric(row.get('Total Hours', 0), errors='coerce')),
                axis=1
            )
        else:
            df['İşveren-Hakediş(USD)'] = pd.to_numeric(df['İşveren Hakediş Birim Fiyat (USD)'], errors='coerce') * pd.to_numeric(df.get('Total Hours', 0), errors='coerce')

    # === 11. İşveren Hakediş (USD) - currency conversion if required ===
    if 'İşveren-Hakediş(USD)' in df.columns and 'Currency' in df.columns and df_info is not None:
        week_col = _find_column(df_info, 'week') or _find_column(df_info, 'month') or (df_info.columns[0] if len(df_info.columns) else None)
        euro_rate_col = _find_column(df_info, 'euro', 'rate') or _find_column(df_info, 'euro')
        def calc_isveren_usd(row):
            currency = str(row.get('Currency', '')).upper()
            amount = pd.to_numeric(row.get('İşveren-Hakediş(USD)', 0), errors='coerce') or 0
            week_month = _normalize_week_label(row.get('(Week / Month)') or row.get('Week / Month') or '')
            try:
                if currency in ('EURO', 'EUR') and week_col and euro_rate_col:
                    exchange_rate = xlookup(week_month, df_info[week_col], df_info[euro_rate_col], 1) or 1
                    return amount * float(exchange_rate)
                else:
                    return amount
            except Exception:
                return amount
        df['İşveren Hakediş (USD)'] = df.apply(calc_isveren_usd, axis=1)

    # === 12. İşveren Hakediş Birim Fiyatı (USD) ===
    if 'İşveren Hakediş (USD)' in df.columns:
        def calc_unit_price(row):
            spec_hours = pd.to_numeric(row.get('Special Hours', 0), errors='coerce')
            denom = spec_hours if (spec_hours and spec_hours > 0) else pd.to_numeric(row.get('Total Hours', 1), errors='coerce')
            denom = denom if denom and denom != 0 else 1
            return pd.to_numeric(row.get('İşveren Hakediş (USD)', 0), errors='coerce') / denom
        df['İşveren-Hakediş Birim Fiyat (USD)'] = df.apply(calc_unit_price, axis=1)

    # === 14-21. Control and lookups from Info sheet ===
    if df_info is not None:
        # Projects lookups
        if 'Projects' in df.columns:
            proj_lookup_col = _find_column(df_info, 'project') or df_info.columns[0]
            control_col = _find_column(df_info, 'control')
            tm_kod_col = _find_column(df_info, 'tm', 'kod') or _find_column(df_info, 'tm')
            kontrol_col = _find_column(df_info, 'kontrol') or _find_column(df_info, 'konrol')
            if proj_lookup_col:
                if control_col:
                    df['Control-1'] = df['Projects'].apply(lambda x: xlookup(x, df_info[proj_lookup_col], df_info[control_col], ''))
                if tm_kod_col:
                    df['TM KOD'] = df['Projects'].apply(lambda x: xlookup(x, df_info[proj_lookup_col], df_info[tm_kod_col], ''))
                if kontrol_col:
                    df['Kontrol-1'] = df['Projects'].apply(lambda x: xlookup(x, df_info[proj_lookup_col], df_info[kontrol_col], ''))

        # Projects/Group lookups for N0-1 etc.
        if 'Projects/Group' in df.columns:
            proj_group_col = _find_column(df_info, 'project', 'group') or df_info.columns[0]
            n0_1 = _find_column(df_info, 'no-1') or _find_column(df_info, 'n0-1') or _find_column(df_info, 'no 1')
            n0_2 = _find_column(df_info, 'no-2') or _find_column(df_info, 'no 2')
            n0_3 = _find_column(df_info, 'no-3') or _find_column(df_info, 'no 3')
            if proj_group_col:
                if n0_1:
                    df['N0-1'] = df['Projects/Group'].apply(lambda x: xlookup(x, df_info[proj_group_col], df_info[n0_1], ''))
                if n0_2:
                    df['NO-2'] = df['Projects/Group'].apply(lambda x: xlookup(x, df_info[proj_group_col], df_info[n0_2], ''))
                if n0_3:
                    df['NO-3'] = df['Projects/Group'].apply(lambda x: xlookup(x, df_info[proj_group_col], df_info[n0_3], ''))

        # TM Liste by Emp
        emp_lookup_col = _find_column(df_info, 'emp') or _find_column(df_info, 'name', 'surname')
        tm_liste_col = _find_column(df_info, 'tm', 'liste') or _find_column(df_info, 'tm')
        if emp_lookup_col and tm_liste_col:
            lookup_field = 'Emp No' if 'Emp No' in df.columns else ('Name Surname' if 'Name Surname' in df.columns else None)
            if lookup_field:
                df['TM Liste'] = df[lookup_field].apply(lambda x: xlookup(x, df_info[emp_lookup_col], df_info[tm_liste_col], ''))

    # === 22. Kontrol-2 --- example compare AN and AO if exist ===
    # If your sheet uses columns 'AN' and 'AO' names, you can map them here; otherwise skip.
    if 'AN' in df.columns and 'AO' in df.columns:
        df['Kontrol-2'] = df['AN'] == df['AO']

    # === KAR/ZARAR calculations ===
    if 'İşveren Hakediş (USD)' in df.columns and 'General Total Cost (USD)' in df.columns:
        df['KAR/ZARAR'] = pd.to_numeric(df['İşveren Hakediş (USD)'], errors='coerce').fillna(0) - pd.to_numeric(df['General Total Cost (USD)'], errors='coerce').fillna(0)
        print(f"Created KAR/ZARAR column with {df['KAR/ZARAR'].notna().sum()} valid values")

    if 'İşveren-Hakediş Birim Fiyat (USD)' in df.columns and 'Hourly Unit Rate (USD)' in df.columns:
        df['BF KAR/ZARAR'] = pd.to_numeric(df['İşveren-Hakediş Birim Fiyat (USD)'], errors='coerce').fillna(0) - pd.to_numeric(df['Hourly Unit Rate (USD)'], errors='coerce').fillna(0)
        print(f"Created BF KAR/ZARAR column with {df['BF KAR/ZARAR'].notna().sum()} valid values")

    return df

def load_favorites():
    """Load favorite reports from JSON file"""
    try:
        with open(FAVORITES_FILE_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if isinstance(data, dict) and '_last_loaded' in data:
                return data
            else:
                return {'reports': data, '_last_loaded': None}
    except FileNotFoundError:
        return {'reports': {}, '_last_loaded': None}

def save_favorites(favorites, last_loaded=None):
    """Save favorite reports to JSON file"""
    data = {
        'reports': favorites,
        '_last_loaded': last_loaded
    }
    with open(FAVORITES_FILE_PATH, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2)

# Routes
@app.route('/')
def home():
    if 'user' not in session:
        return redirect('/login.html')
    return redirect('/index.html')

@app.route('/index.html')
@login_required
def index():
    """Main dashboard page"""
    return render_template('index.html', user=session.get('user'), role=session.get('role'))

@app.route('/table.html')
@login_required
def table():
    """Pivot analysis page"""
    return render_template('table.html', user=session.get('user'), role=session.get('role'))

@app.route('/profile.html')
@login_required
def profile():
    """User profile page"""
    user = User.query.filter_by(username=session.get('user')).first()
    return render_template('profile.html', 
                         user=session.get('user'), 
                         name=user.name if user else session.get('user'),
                         role=session.get('role'))

@app.route('/login.html')
def login_page():
    """Login page"""
    return render_template('login.html')

@app.route('/register.html')
def register_page():
    """Registration page"""
    return render_template('register.html')

@app.route('/graphs.html')
@login_required
def graphs():
    """Graph analysis page"""
    return render_template('graphs.html', user=session.get('user'), role=session.get('role'))

@app.route('/admin.html')
@login_required
def admin_panel():
    """Admin panel - Admin only"""
    if session.get('role') != 'admin':
        return redirect('/index.html')
    return render_template('admin.html', user=session.get('user'), role=session.get('role'), name=session.get('name'))

# Simple rate limiting storage (in production, use Redis)
_login_attempts = {}

@app.route('/api/login', methods=['POST'])
def login():
    """Handle user login with rate limiting"""
    data = request.json
    username = data.get('username', '').strip()
    password = data.get('password', '')
    
    if not username or not password:
        return jsonify({'error': 'Username and password required'}), 400
    
    # Rate limiting check
    client_ip = request.remote_addr
    current_time = datetime.now().timestamp()
    
    if client_ip in _login_attempts:
        attempts, last_attempt = _login_attempts[client_ip]
        if current_time - last_attempt < 60:  # Within 1 minute
            if attempts >= 5:
                return jsonify({'error': 'Too many login attempts. Please try again in 1 minute.'}), 429
        else:
            _login_attempts[client_ip] = (0, current_time)
    
    # Check if user exists in database
    user = User.query.filter_by(username=username).first()
    if user and check_password_hash(user.password, password):
        # Reset failed attempts on successful login
        _login_attempts.pop(client_ip, None)
        
        session['user'] = username
        session['role'] = user.role
        session['name'] = user.name
        session.permanent = True
        return jsonify({
            'success': True,
            'role': user.role,
            'name': user.name
        })
    
    # Track failed attempts
    if client_ip in _login_attempts:
        attempts, _ = _login_attempts[client_ip]
        _login_attempts[client_ip] = (attempts + 1, current_time)
    else:
        _login_attempts[client_ip] = (1, current_time)
    
    return jsonify({'error': 'Invalid username or password'}), 401

@app.route('/api/register', methods=['POST'])
def register():
    """Handle user registration"""
    data = request.json
    username = data.get('username', '').strip()
    password = data.get('password', '')
    name = data.get('name', '').strip()
    
    if not username or not password or not name:
        return jsonify({'error': 'All fields required'}), 400
    
    # Password strength validation
    if len(password) < 8:
        return jsonify({'error': 'Password must be at least 8 characters'}), 400
    if not any(c.isupper() for c in password):
        return jsonify({'error': 'Password must contain at least one uppercase letter'}), 400
    if not any(c.isdigit() for c in password):
        return jsonify({'error': 'Password must contain at least one number'}), 400
    
    # Username validation
    if len(username) < 3:
        return jsonify({'error': 'Username must be at least 3 characters'}), 400
    if not username.isalnum():
        return jsonify({'error': 'Username must contain only letters and numbers'}), 400
    
    # Check if username exists
    if User.query.filter_by(username=username).first():
        return jsonify({'error': 'Username already exists'}), 400
    
    # Register new user (normal user role)
    new_user = User(
        username=username,
        password=generate_password_hash(password),
        name=name,
        role='user'
    )
    db.session.add(new_user)
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Registration successful'})

@app.route('/api/profile', methods=['GET'])
@login_required
def get_profile():
    """Get current user profile"""
    user = User.query.filter_by(username=session.get('user')).first()
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    return jsonify({
        'username': user.username,
        'name': user.name,
        'email': user.email or '',
        'first_name': user.first_name or '',
        'last_name': user.last_name or '',
        'profile_photo': user.profile_photo,
        'role': user.role
    })

@app.route('/api/profile', methods=['PUT'])
@login_required
def update_profile():
    """Update user profile"""
    try:
        user = User.query.filter_by(username=session.get('user')).first()
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        data = request.json
        
        # Update fields if provided
        if 'email' in data:
            user.email = data['email']
        if 'first_name' in data:
            user.first_name = data['first_name']
        if 'last_name' in data:
            user.last_name = data['last_name']
        if 'name' in data:
            user.name = data['name']
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Profile updated successfully'
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/profile/photo', methods=['POST'])
@login_required
def upload_profile_photo():
    """Upload profile photo"""
    try:
        if 'photo' not in request.files:
            return jsonify({'error': 'No photo provided'}), 400
        
        file = request.files['photo']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Check file extension
        allowed_extensions = {'png', 'jpg', 'jpeg', 'gif'}
        file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
        if file_ext not in allowed_extensions:
            return jsonify({'error': 'Invalid file type. Use PNG, JPG, or GIF'}), 400
        
        user = User.query.filter_by(username=session.get('user')).first()
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        # Create profile photos directory if it doesn't exist
        photo_dir = os.path.join('static', 'img', 'profiles')
        os.makedirs(photo_dir, exist_ok=True)
        
        # Save photo with user ID in filename
        filename = f'user_{user.id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.{file_ext}'
        filepath = os.path.join(photo_dir, filename)
        file.save(filepath)
        
        # Update user profile photo path
        user.profile_photo = f'img/profiles/{filename}'
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Photo uploaded successfully',
            'photo_url': f'/static/{user.profile_photo}'
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/logout', methods=['POST'])
def logout():
    """Handle user logout"""
    session.clear()
    return jsonify({'success': True})

# List of calculated fields that should not be in add record form
CALCULATED_FIELDS = [
    'North/South', 'North/ South', 'Currency', 'Hourly Rate', 'Cost', 'Hourly Base Rate',
    'Hourly Additional Rate', 'Hourly Additional Rates', 'AP-CB/Subcon', 'AP-CB / Subcon',
    'General Total Cost (USD)', 'Hourly Unit Rate (USD)', 
    'İşveren Hakediş Birim Fiyat', 'İşveren Hakediş Birim Fiyat (USD)', 'İşveren-Hakediş Birim Fiyat',
    'İşveren-Hakediş(USD)', 'İşveren Hakediş (USD)', 'İşveren-Hakediş Birim Fiyat (USD)',
    'İşveren- Hakediş (USD)', 'İşveren-Hakediş (USD)',
    'İsveren - Currency', 'İşveren - Currency',
    'İşveren- Sözleşme No', 'İşveren Sözleşme No', 'İşveren-Sözleşme No',
    'İşveren-Hakediş Kapsam', 'İşveren Hakediş Kapsam', 'İşveren- Hakediş Kapsam',
    'İşveren-Hakediş Dönemi', 'İşveren Hakediş Dönemi', 'İşveren- Hakediş Dönemi',
    'İşveren-Hakediş No', 'İşveren Hakediş No', 'İşveren- Hakediş No',
    'İşveren- Hakediş', 'İşveren-Hakediş', 'İşveren Hakediş',
    'İşveren- MH-Modifiye', 'İşveren MH-Modifiye', 'İşveren-MH-Modifiye',
    'İşveren- Actual (USD)', 'İşveren Actual (USD)', 'İşveren-Actual (USD)',
    'Control-1', 'TM Liste', 'TM KOD', 'TM Kod', 'N0-1', 'NO-1', 'Kontrol-1', 'Konrol-1', 'Kontrol-2', 'Knrtol-2',
    'NO-2', 'NO-3', 'NO-10', 'LS/Unit Rate', 'KAR/ZARAR', 'BF KAR/ZARAR'
]

@app.route('/api/add-record', methods=['POST'])
@login_required
def add_record():
    """Add new database record - Admin only"""
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    try:
        data = request.json
        record_data = data.get('record', {})
        
        # Get personel from either PERSONEL or Name Surname field
        personel = record_data.get('PERSONEL', '') or record_data.get('Name Surname', '')
        
        if not personel:
            return jsonify({'error': 'Name Surname (PERSONEL) field is required'}), 400
        
        # Ensure both fields are set for consistency
        if 'Name Surname' in record_data and 'PERSONEL' not in record_data:
            record_data['PERSONEL'] = record_data['Name Surname']
        
        # Create DataFrame with single row to calculate formulas
        df_single = pd.DataFrame([record_data])
        df_single = add_calculated_columns(df_single)
        
        # Get the calculated record with all formulas applied
        calculated_record = df_single.iloc[0].to_dict()
        
        # Convert numpy types to Python types for JSON serialization
        for key, value in calculated_record.items():
            if pd.isna(value):
                calculated_record[key] = None
            elif hasattr(value, 'item'):  # numpy types
                calculated_record[key] = value.item()
        
        # Create new record with calculated fields
        new_record = DatabaseRecord(
            personel=personel,
            data=json.dumps(calculated_record)
        )
        db.session.add(new_record)
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Record added successfully'})
    except Exception as e:
        db.session.rollback()
        print(f"Add record error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/get-records', methods=['GET'])
@login_required
def get_records():
    """Get database records filtered by user"""
    try:
        # Pagination parameters
        try:
            page = int(request.args.get('page', 1))
            per_page = int(request.args.get('per_page', 10))
        except:
            page = 1
            per_page = 10

        # Search parameter (search by personel / Name Surname)
        search = (request.args.get('search') or '').strip()

        # Build base query depending on role
        if session.get('role') == 'admin':
            base_query = DatabaseRecord.query
        else:
            base_query = DatabaseRecord.query.filter_by(personel=session.get('name'))

        # Apply search filter on the personel column
        if search:
            try:
                base_query = base_query.filter(DatabaseRecord.personel.ilike(f"%{search}%"))
            except Exception:
                # Fallback to simple filter (case-sensitive) if ilike not supported
                base_query = base_query.filter(DatabaseRecord.personel.contains(search))

        # Use SQLAlchemy pagination to avoid loading everything into memory
        pagination = base_query.order_by(DatabaseRecord.id.desc()).paginate(page=page, per_page=per_page, error_out=False)
        items = pagination.items

        records_list = []
        for record in items:
            record_dict = json.loads(record.data)
            record_dict['id'] = record.id
            records_list.append(record_dict)

        return jsonify({
            'success': True,
            'records': records_list,
            'page': page,
            'per_page': per_page,
            'pages': pagination.pages,
            'total': pagination.total
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/get-record/<int:record_id>', methods=['GET'])
@login_required
def get_record(record_id):
    """Get single database record by id"""
    try:
        record = DatabaseRecord.query.get(record_id)
        if not record:
            return jsonify({'error': 'Record not found'}), 404

        # Security: non-admins may only see their own records
        if session.get('role') != 'admin' and record.personel != session.get('name'):
            return jsonify({'error': 'Access denied'}), 403

        record_dict = json.loads(record.data)
        record_dict['id'] = record.id
        return jsonify({'success': True, 'record': record_dict})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/update-record/<int:record_id>', methods=['PUT'])
@login_required
def update_record(record_id):
    """Update database record - Admin only"""
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    try:
        record = DatabaseRecord.query.get(record_id)
        if not record:
            return jsonify({'error': 'Record not found'}), 404
        
        data = request.json
        record_data = data.get('record', {})
        
        record.personel = record_data.get('PERSONEL', record.personel)
        record.data = json.dumps(record_data)
        record.updated_at = datetime.utcnow()
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'Record updated successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

def is_calculated_field(field_name):
    """Check if a field is calculated based on patterns"""
    field_lower = str(field_name).lower()
    
    # Check against explicit list
    if field_name in CALCULATED_FIELDS:
        return True
    
    # Pattern-based detection
    calculated_patterns = [
        'kar/zarar', 'kar zarar',
        'hourly rate', 'hourly base', 'hourly additional', 'hourly unit',
        'cost (usd)', 'total cost',
        'hakediş', 'hakedi',
        'currency', 'north/south', 'north/ south',
        'control-', 'kontrol-', 'konrol-', 'knrtol-',
        'tm kod', 'tm liste',
        'no-1', 'no-2', 'no-3', 'no-10',
        'ls/unit rate',
        'ap-cb/subcon', 'ap-cb / subcon'
    ]
    
    for pattern in calculated_patterns:
        if pattern in field_lower:
            return True
    
    return False

@app.route('/api/get-input-fields', methods=['GET'])
@login_required
def get_input_fields():
    """Get list of fields for add record form (excluding calculated fields)"""
    try:
        # Get latest file to determine available columns
        file_path = session.get('current_file')
        if not file_path or not os.path.exists(file_path):
            upload_dir = app.config['UPLOAD_FOLDER']
            xlsb_files = [f for f in os.listdir(upload_dir) if f.endswith('.xlsb')]
            if not xlsb_files:
                return jsonify({'error': 'No data file found'}), 404
            xlsb_files.sort(reverse=True)
            file_path = os.path.join(upload_dir, xlsb_files[0])
        
        # Load first row to get column names
        df = load_excel_data(file_path)
        
        # Get all columns except calculated ones
        input_fields = [col for col in df.columns if not is_calculated_field(col)]
        calculated_fields = [col for col in df.columns if is_calculated_field(col)]
        
        print(f"Total columns: {len(df.columns)}")
        print(f"Input fields: {len(input_fields)}")
        print(f"Calculated fields: {len(calculated_fields)}")
        print(f"Calculated fields list: {calculated_fields[:10]}...")  # Print first 10
        
        return jsonify({
            'success': True,
            'input_fields': input_fields,
            'calculated_fields': calculated_fields
        })
    except Exception as e:
        print(f"Get input fields error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/delete-record/<int:record_id>', methods=['DELETE'])
@login_required
def delete_record(record_id):
    """Delete database record - Admin only"""
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    try:
        record = DatabaseRecord.query.get(record_id)
        if not record:
            return jsonify({'error': 'Record not found'}), 404
        
        db.session.delete(record)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Record deleted successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/get-person-info', methods=['GET'])
@login_required
def get_person_info():
    """Get person information from Info and Hourly Rates sheets"""
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    try:
        name = request.args.get('name', '').strip()
        if not name:
            return jsonify({'error': 'Name parameter required'}), 400
        
        # Get latest uploaded file
        file_path = session.get('current_file')
        if not file_path or not os.path.exists(file_path):
            # Try to find latest xlsb file in uploads
            upload_dir = app.config['UPLOAD_FOLDER']
            xlsb_files = [f for f in os.listdir(upload_dir) if f.endswith('.xlsb')]
            if not xlsb_files:
                return jsonify({'error': 'No Excel file found'}), 404
            xlsb_files.sort(reverse=True)
            file_path = os.path.join(upload_dir, xlsb_files[0])
        
        # Read Info sheet
        try:
            df_info = pd.read_excel(file_path, sheet_name='Info', engine='pyxlsb')
        except Exception as e:
            return jsonify({'error': f'Could not read Info sheet: {str(e)}'}), 500
        
        # Read Hourly Rates sheet (header is on row 2)
        try:
            df_rates = pd.read_excel(file_path, sheet_name='Hourly Rates', engine='pyxlsb', header=1)
        except Exception as e:
            return jsonify({'error': f'Could not read Hourly Rates sheet: {str(e)}'}), 500
        
        # Search in Info sheet by Name column
        info_match = df_info[df_info['Name'].str.strip().str.upper() == name.upper()]
        
        # Search in Hourly Rates by Name Surname column
        rates_match = df_rates[df_rates['Name Surname'].str.strip().str.upper() == name.upper()]
        
        result = {}
        
        # Extract data from Info sheet
        if not info_match.empty:
            row = info_match.iloc[0]
            result['Company'] = str(row.get('Company', '')) if pd.notna(row.get('Company')) else ''
            result['Nationality'] = str(row.get('Nationality', '')) if pd.notna(row.get('Nationality')) else ''
            result['Discipline'] = str(row.get('Discipline', '')) if pd.notna(row.get('Discipline')) else ''
            result['Scope'] = str(row.get('Scope', '')) if pd.notna(row.get('Scope')) else ''
            result['Projects'] = str(row.get('Projects', '')) if pd.notna(row.get('Projects')) else ''
            result['Projects/Group'] = str(row.get('Projects/Group', '')) if pd.notna(row.get('Projects/Group')) else ''
            result['North/South'] = str(row.get('North/South', '')) if pd.notna(row.get('North/South')) else ''
        
        # Extract data from Hourly Rates sheet
        if not rates_match.empty:
            row = rates_match.iloc[0]
            # Get hourly rates (trying different columns)
            for col in ['Hourly Base Rates 1', 'Hourly Base Rates 2', 'Hourly Base Rates 3']:
                if col in row and pd.notna(row[col]) and row[col] != 0:
                    result['Hourly Rate'] = float(row[col])
                    break
        
        if not result:
            return jsonify({'error': 'Person not found in Info or Hourly Rates sheets'}), 404
        
        return jsonify({
            'success': True,
            'data': result
        })
    
    except Exception as e:
        print(f"Get person info error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/clear-database', methods=['POST'])
@login_required
def clear_database():
    """Clear all database records - Admin only"""
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    try:
        import sys
        print("\n" + "="*80, file=sys.stderr, flush=True)
        print("DEBUG: Starting database clear...", file=sys.stderr, flush=True)
        
        deleted = DatabaseRecord.query.delete()
        print(f"DEBUG: Deleted {deleted} records", file=sys.stderr, flush=True)
        
        db.session.commit()
        print("DEBUG: Database commit successful", file=sys.stderr, flush=True)
        
        # Clear cache
        clear_data_cache()
        
        # Clear session file reference
        session.pop('current_file', None)
        session.pop('data_shape', None)
        
        print("DEBUG: Clear database completed successfully", file=sys.stderr, flush=True)
        print("="*80 + "\n", file=sys.stderr, flush=True)
        
        return jsonify({
            'success': True, 
            'deleted': deleted,
            'message': f'Database cleared successfully! {deleted} records deleted.'
        })
    except Exception as e:
        import sys
        import traceback
        print("\n" + "="*80, file=sys.stderr, flush=True)
        print(f"ERROR clearing database: {str(e)}", file=sys.stderr, flush=True)
        traceback.print_exc(file=sys.stderr)
        print("="*80 + "\n", file=sys.stderr, flush=True)
        
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/forgot-password.html')
def forgot_password():
    """Forgot password page"""
    return render_template('forgot-password.html')

# API Routes
@app.route('/api/check-session', methods=['GET'])
@login_required
def check_session():
    """Check if file is loaded in session or data exists in database"""
    try:
        # Get user filter (None for admin, name for regular users)
        user_filter = None if session.get('role') == 'admin' else session.get('name')
        
        # Get combined data from file and database
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        
        if df.empty:
            return jsonify({'hasData': False})
        
        # Add calculated columns (will use cached if available)
        df_with_calc = add_calculated_columns(df)
        
        # Get all columns including calculated ones
        all_columns = df_with_calc.columns.tolist()
        
        # Remove PERSONEL from column list if Name Surname exists (it's internal only)
        if 'Name Surname' in all_columns and 'PERSONEL' in all_columns:
            all_columns.remove('PERSONEL')
        
        # Limit preview to first 5 rows to prevent page collapse
        df_preview = df_with_calc.head(5).fillna('')
        data_json = df_preview.to_dict('records')
        
        # Get filter columns for ALL categorical columns
        filter_cols = []
        # Skip these columns from filters (PERSONEL is internal, but Name Surname is allowed)
        skip_cols = ['PERSONEL', 'id', 'created_at', 'updated_at']
        
        # Always add Name Surname filter first (important for filtering by person)
        if 'Name Surname' in df_with_calc.columns:
            name_values = sorted([str(v) for v in df_with_calc['Name Surname'].dropna().unique()])
            if name_values:
                filter_cols.append({
                    'name': 'Name Surname',
                    'values': name_values
                })
        
        for col in df_with_calc.columns:
            # Skip numeric columns and internal columns
            if col in skip_cols or col == 'Name Surname':  # Skip Name Surname as it's already added
                continue
                
            # Check if column is categorical (not purely numeric)
            try:
                # Try to detect categorical columns
                unique_count = df_with_calc[col].nunique()
                total_count = len(df_with_calc[col].dropna())
                
                # If column has reasonable number of unique values, add to filters
                # Max 200 unique values to accommodate other columns
                if unique_count > 0 and unique_count <= 200:
                    # Check if it's not a purely numeric column with many values
                    non_null_values = df_with_calc[col].dropna()
                    if len(non_null_values) > 0:
                        # Debug BEFORE creating unique values
                        if 'week' in col.lower() or 'month' in col.lower():
                            import sys
                            print("\n" + "="*80, file=sys.stderr, flush=True)
                            print(f"DEBUG RAW DF: Column '{col}'", file=sys.stderr, flush=True)
                            print(f"DEBUG RAW DF: First 10 raw values from df: {df_with_calc[col].head(10).tolist()}", file=sys.stderr, flush=True)
                            print(f"DEBUG RAW DF: Dtype: {df_with_calc[col].dtype}", file=sys.stderr, flush=True)
                        
                        # Add to filters
                        unique_values = sorted([str(v) for v in df_with_calc[col].dropna().unique()])
                        
                        # Debug date column - FORCE OUTPUT
                        if 'week' in col.lower() or 'month' in col.lower():
                            print(f"DEBUG FILTER: Total unique values: {len(unique_values)}", file=sys.stderr, flush=True)
                            print(f"DEBUG FILTER: First 5: {unique_values[:5]}", file=sys.stderr, flush=True)
                            print(f"DEBUG FILTER: Last 5: {unique_values[-5:]}", file=sys.stderr, flush=True)
                            if unique_values:
                                print(f"DEBUG FILTER: Sample length: {len(unique_values[0])}", file=sys.stderr, flush=True)
                            print("="*80 + "\n", file=sys.stderr, flush=True)
                        
                        filter_cols.append({
                            'name': col,
                            'values': unique_values
                        })
            except:
                continue
        
        return jsonify({
            'hasData': True,
            'columns': all_columns,
            'shape': df_with_calc.shape,
            'data': data_json,  # Preview only (first 5 rows)
            'filter_columns': filter_cols
        })
    except Exception as e:
        print(f"Check session error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'hasData': False, 'error': str(e)})

@app.route('/api/upload', methods=['POST'])
@login_required
def upload_file():
    """Handle file upload and process Excel data"""
    try:
        # Clear cache on new upload
        clear_data_cache()
        
        # Only admin can upload files
        if session.get('role') != 'admin':
            return jsonify({'error': 'Only admin can upload files'}), 403
        
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not file:
            return jsonify({'error': 'Invalid file'}), 400
            
        # Check file extension
        allowed_extensions = {'.xlsx', '.xls', '.xlsb'}
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in allowed_extensions:
            return jsonify({'error': f'Invalid file type. Allowed: {", ".join(allowed_extensions)}'}), 400
        
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'{timestamp}_{filename}'
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        # Save file
        print(f"Saving file to: {filepath}")
        file.save(filepath)
        print(f"File saved successfully")
        
        # Load and process data quickly (don't apply calculated columns yet for speed)
        print(f"Loading Excel data...")
        df = load_excel_data(filepath)
        print(f"Data loaded: {df.shape[0]} rows, {df.shape[1]} columns")
        
        # Save data to database permanently
        print(f"Saving data to database...")
        saved_count = 0
        skipped_count = 0
        
        # Clear existing database records (optional - remove if you want to keep old data)
        # DatabaseRecord.query.delete()
        
        for idx, row in df.iterrows():
            try:
                # Convert row to dictionary
                row_dict = row.to_dict()
                
                # Get personel name from either PERSONEL or Name Surname column
                personel = row_dict.get('PERSONEL') or row_dict.get('Name Surname', '')
                
                if not personel:
                    skipped_count += 1
                    continue
                
                # Ensure both fields exist for consistency
                if 'Name Surname' in row_dict and 'PERSONEL' not in row_dict:
                    row_dict['PERSONEL'] = row_dict['Name Surname']
                elif 'PERSONEL' in row_dict and 'Name Surname' not in row_dict:
                    row_dict['Name Surname'] = row_dict['PERSONEL']
                
                # Convert any NaN values to empty strings
                for key, value in row_dict.items():
                    if pd.isna(value):
                        row_dict[key] = ''
                
                # Debug: Check date format before saving to database
                if '(Week / Month)' in row_dict:
                    if saved_count == 0:  # Only log first record
                        print(f"DEBUG SAVE: First date value before DB save: {row_dict['(Week / Month)']}")
                
                # Create database record
                new_record = DatabaseRecord(
                    personel=str(personel),
                    data=json.dumps(row_dict)
                )
                db.session.add(new_record)
                saved_count += 1
                
                # Commit in batches for better performance
                if saved_count % 100 == 0:
                    db.session.commit()
                    print(f"Saved {saved_count} records...")
                    
            except Exception as e:
                print(f"Error saving row {idx}: {str(e)}")
                skipped_count += 1
                continue
        
        # Final commit
        db.session.commit()
        print(f"Database save complete: {saved_count} saved, {skipped_count} skipped")
        
        # Store in session
        session['current_file'] = filepath
        session['data_shape'] = df.shape
        
        # Get important filter columns only (for speed)
        filter_cols = []
        important_cols = ['PERSONEL', 'Name Surname', 'Company', 'Projects', 'Status', 'Discipline']
        for col in important_cols:
            if col in df.columns:
                unique_values = df[col].dropna().unique()
                if len(unique_values) <= 50:
                    filter_cols.append({
                        'name': col,
                        'values': [str(v) for v in unique_values[:50]]
                    })
        
        # Apply calculated columns after basic info is ready
        df = add_calculated_columns(df)
        
        # Convert to JSON-friendly format (full dataset)
        df_clean = df.fillna('')
        data_json = df_clean.to_dict('records')
        
        print(f"Upload successful, returning response")
        return jsonify({
            'success': True,
            'shape': df.shape,
            'columns': df.columns.tolist(),
            'filter_columns': filter_cols,
            'data': data_json,
            'saved_to_db': saved_count,
            'skipped': skipped_count,
            'message': f'File uploaded successfully! {df.shape[0]} rows, {df.shape[1]} columns. Saved {saved_count} records to database.'
        })
    
    except Exception as e:
        print(f"Upload error: {str(e)}")
        import traceback
        traceback.print_exc()
        db.session.rollback()  # Rollback any pending transactions
        return jsonify({'error': str(e)}), 500

@app.route('/api/filter', methods=['POST'])
@login_required
def filter_data():
    """Apply filters to data"""
    try:
        filters = request.json.get('filters', {})
        
        # Get user filter
        user_filter = None if session.get('role') == 'admin' else session.get('name')
        
        # Load combined data
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        df = add_calculated_columns(df)
        
        # Apply filters
        for col, values in filters.items():
            if col in df.columns and values:
                # Convert both to strings for consistent comparison
                df_col_str = df[col].astype(str)
                values_str = [str(v) for v in values]
                df = df[df_col_str.isin(values_str)]
        
        # Convert to JSON (replace NaN with empty string)
        df_clean = df.fillna('')
        data_json = df_clean.to_dict('records')
        
        return jsonify({
            'success': True,
            'data': data_json,
            'shape': df.shape
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/get-filtered-options', methods=['POST'])
@login_required
def get_filtered_options():
    """Get available filter options based on current filter selections (for cascading filters)"""
    try:
        filters = request.json.get('filters', {})
        
        # Get user filter
        user_filter = None if session.get('role') == 'admin' else session.get('name')
        
        # Load combined data
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        df = add_calculated_columns(df)
        
        # Apply filters
        for col, values in filters.items():
            if col in df.columns and values:
                # Convert both to strings for consistent comparison
                df_col_str = df[col].astype(str)
                values_str = [str(v) for v in values]
                df = df[df_col_str.isin(values_str)]
        
        # Get available options for each column after filtering
        filter_cols = []
        skip_cols = ['PERSONEL', 'id', 'created_at', 'updated_at']
        
        # Always add Name Surname filter first
        if 'Name Surname' in df.columns:
            name_values = sorted([str(v) for v in df['Name Surname'].dropna().unique()])
            if name_values:
                filter_cols.append({
                    'name': 'Name Surname',
                    'values': name_values
                })
        
        for col in df.columns:
            if col in skip_cols or col == 'Name Surname':  # Skip Name Surname as it's already added
                continue
                
            try:
                unique_count = df[col].nunique()
                
                if unique_count > 0 and unique_count <= 200:
                    non_null_values = df[col].dropna()
                    if len(non_null_values) > 0:
                        unique_values = sorted([str(v) for v in df[col].dropna().unique()])
                        
                        filter_cols.append({
                            'name': col,
                            'values': unique_values
                        })
            except:
                continue
        
        return jsonify({
            'success': True,
            'filter_columns': filter_cols
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/pivot', methods=['POST'])
@login_required
def create_pivot():
    """Create pivot table"""
    try:
        config = request.json
        index_col = config.get('index')
        columns_col = config.get('columns')
        values_cols = config.get('values', [])
        agg_func = config.get('agg_func', 'sum')
        filters = config.get('filters', {})
        
        print(f"\n{'='*80}")
        print(f"PIVOT REQUEST DEBUG")
        print(f"{'='*80}")
        print(f"Index: {index_col}")
        print(f"Columns: {columns_col}")
        print(f"Values: {values_cols}")
        print(f"Agg Function: {agg_func}")
        print(f"Filters received: {len(filters)} categories")
        for col, vals in filters.items():
            print(f"  - {col}: {len(vals)} values")
        print(f"{'='*80}\n")
        
        # Get user filter
        user_filter = None if session.get('role') == 'admin' else session.get('name')
        
        # Load and filter data
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        df = add_calculated_columns(df)
        
        print(f"Data shape before filters: {df.shape}")
        
        # Apply filters
        for col, values in filters.items():
            if col in df.columns and values:
                before_count = len(df)
                # Convert both dataframe column and filter values to strings for comparison
                # This ensures consistent comparison regardless of data type
                df_col_str = df[col].astype(str)
                values_str = [str(v) for v in values]
                df = df[df_col_str.isin(values_str)]
                after_count = len(df)
                print(f"Filter '{col}': {before_count} → {after_count} rows (filtered out {before_count - after_count})")
                # Debug: Show what we're filtering
                if after_count == 0:
                    print(f"  WARNING: All data filtered out for column '{col}'")
                    print(f"  Filter values: {values_str[:10]}")
                    print(f"  Unique values in data: {df_col_str.unique()[:10].tolist()}")
        
        print(f"Data shape after filtering: {df.shape}")
        print(f"{'='*80}\n")
        # If no rows remain after filtering, return a clear error
        if df.empty:
            return jsonify({'error': 'No data available after applying filters. Please relax filters or select a different dataset.'}), 400
        
        # Create pivot
        if index_col and values_cols:
            # Validate that index column exists
            if index_col not in df.columns:
                return jsonify({'error': f'Index column "{index_col}" not found in data'}), 400
            
            # Validate and convert value columns to numeric
            valid_values = []
            for col in values_cols:
                print(f"\n--- Processing value column: '{col}' ---")
                
                if col not in df.columns:
                    print(f"❌ Column '{col}' not found in dataframe")
                    print(f"Available columns: {df.columns.tolist()}")
                    continue
                
                # Debug: Show sample values
                print(f"Column type: {df[col].dtype}")
                print(f"First 5 values: {df[col].head(5).tolist()}")
                print(f"Non-null count: {df[col].notna().sum()} / {len(df)}")
                    
                # Try to convert to numeric and ensure it's a Series
                try:
                    # Attempt robust numeric conversion: strip common currency/formatting
                    series = df[col]
                    # If object/string dtype, clean common characters like commas, currency symbols and words
                    if series.dtype == object or str(series.dtype).startswith('string'):
                        s_clean = series.astype(str)
                        # Remove common currency symbols and abbreviations, keep digits, dot and minus
                        s_clean = s_clean.str.replace(r'[\$,£€¥]', '', regex=True)
                        s_clean = s_clean.str.replace(r'\bUSD\b|\bTRY\b|\bEUR\b|\bTL\b', '', regex=True, case=False)
                        # Remove any non-numeric characters except dot and minus
                        s_clean = s_clean.str.replace(r'[^0-9.\-]', '', regex=True)
                        numeric_series = pd.to_numeric(s_clean, errors='coerce')
                    else:
                        numeric_series = pd.to_numeric(series, errors='coerce')
                    print(f"After cleaning + pd.to_numeric - type: {type(numeric_series)}, dtype: {numeric_series.dtype}")
                    
                    # Ensure it's 1-dimensional
                    if hasattr(numeric_series, 'ndim'):
                        print(f"Dimension check: ndim = {numeric_series.ndim}")
                        if numeric_series.ndim != 1:
                            print(f"❌ Column is {numeric_series.ndim}-dimensional, skipping")
                            continue
                    
                    # Check if we have any valid numeric values
                    valid_count = numeric_series.notna().sum()
                    null_count = numeric_series.isna().sum()
                    print(f"Valid numeric values: {valid_count}")
                    print(f"Null/NaN values: {null_count}")
                    print(f"Sample numeric values: {numeric_series.dropna().head(5).tolist()}")
                    
                    if valid_count == 0:
                        print(f"❌ No valid numeric values found")
                        continue
                    
                    # Update the dataframe with numeric values
                    df[col] = numeric_series
                    valid_values.append(col)
                    print(f"✅ Successfully converted '{col}' to numeric: {valid_count} valid values")
                except Exception as e:
                    print(f"❌ Error processing column '{col}': {str(e)}")
                    import traceback
                    traceback.print_exc()
                    continue
            
            print(f"\n{'='*80}")
            print(f"VALID COLUMNS: {valid_values}")
            print(f"{'='*80}\n")
            
            if not valid_values:
                return jsonify({'error': 'No valid numeric columns selected for analysis. Please select columns with numeric values (costs, rates, etc.)'}), 400
            
            # Build pivot parameters
            pivot_params = {
                'values': valid_values[0] if len(valid_values) == 1 else valid_values,
                'index': index_col,
                'aggfunc': agg_func,
                'fill_value': 0  # Fill NaN with 0
            }
            
            # Add columns parameter if specified
            if columns_col and columns_col in df.columns:
                pivot_params['columns'] = columns_col
            
            print(f"Creating pivot with params: {pivot_params}")
            
            # Create pivot table
            pivot = pd.pivot_table(df, **pivot_params)
            pivot = pivot.reset_index()
            
            print(f"Pivot created successfully: {pivot.shape}")
            
            # Replace NaN with 0 for display
            pivot = pivot.fillna(0)
            
            # Format column names if they're tuples (multi-level columns)
            if any(isinstance(col, tuple) for col in pivot.columns):
                pivot.columns = [' - '.join(map(str, col)).strip(' -') if isinstance(col, tuple) else str(col) for col in pivot.columns]
            
            return jsonify({
                'success': True,
                'data': pivot.to_dict('records'),
                'columns': pivot.columns.tolist()
            })
        else:
            return jsonify({'error': 'Please select both Group By column and at least one Value column'}), 400
    
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"Pivot error: {str(e)}")
        print(error_trace)
        return jsonify({'error': f'Error creating pivot table: {str(e)}'}), 500

@app.route('/api/chart', methods=['POST'])
@login_required
def create_chart():
    """Create chart from data"""
    try:
        config = request.json
        chart_type = config.get('chart_type')
        # Support both old and new parameter names
        x_col = config.get('x_column') or config.get('x')
        y_col = config.get('y_column') or config.get('y')
        color_col = config.get('color_column') or config.get('color')
        filters = config.get('filters', {})
        
        # Always use whole database for graphs (no user filter)
        user_filter = None
        
        # Load and filter data
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        
        if df.empty:
            return jsonify({'error': 'No data available'}), 400
        
        # Apply filters BEFORE adding calculated columns for better performance
        if filters:
            for col, values in filters.items():
                if col in df.columns and values:
                    df = df[df[col].isin(values)]
        
        # Add calculated columns to ALL data (don't limit yet)
        df = add_calculated_columns(df)
        
        # Validate columns exist
        if x_col and x_col not in df.columns:
            return jsonify({'error': f'Column "{x_col}" not found'}), 400
        if y_col and y_col not in df.columns:
            return jsonify({'error': f'Column "{y_col}" not found'}), 400
        
        # Clean data for plotting - handle mixed types
        if x_col:
            df[x_col] = df[x_col].fillna('')
            # No limit on unique values - show all data
        
        if y_col:
            # Convert y column to numeric if possible
            df[y_col] = pd.to_numeric(df[y_col], errors='coerce').fillna(0)
            # Remove rows with zero or negative values for pie charts
            if chart_type == 'pie':
                df = df[df[y_col] > 0]
        
        if color_col and color_col in df.columns:
            df[color_col] = df[color_col].fillna('')
        
        # Remove rows with invalid data
        df = df.dropna(subset=[col for col in [x_col, y_col] if col and col in df.columns])
        
        if df.empty:
            return jsonify({'error': 'No valid data after filtering'}), 400
        
        # For line charts: prefer date format over week codes if both exist
        if chart_type == 'line' and x_col and x_col in df.columns:
            import re
            # Check if we have date format entries (containing / and NOT matching W##)
            df[x_col] = df[x_col].astype(str)
            has_dates = df[x_col].str.contains('/', na=False).any()
            has_week_codes = df[x_col].str.match(r'^W\d+$', case=False, na=False).any()
            
            # Only filter out week codes if we also have date format entries
            if has_dates and has_week_codes:
                # Keep only rows that contain "/" (date format)
                df = df[df[x_col].str.contains('/', na=False)]
        
        if df.empty:
            return jsonify({'error': 'No valid data after filtering'}), 400
        
        # Create chart
        fig = None
        color_param = color_col if color_col and color_col in df.columns else None
        
        # Sort data by X column if it looks like a date (for proper ordering)
        x_sort_col = None
        if x_col:
            try:
                # Create a temporary datetime column for sorting only
                x_sort_col = f'{x_col}_sort_temp'
                import warnings
                with warnings.catch_warnings():
                    warnings.filterwarnings('ignore', category=UserWarning)
                    df[x_sort_col] = pd.to_datetime(df[x_col], errors='coerce', format='mixed')
                if df[x_sort_col].notna().any():
                    df = df.sort_values(x_sort_col)
                    df = df.drop(columns=[x_sort_col])
            except:
                pass
        
        # Use simpler chart types for better performance
        if chart_type == 'bar':
            # Aggregate data by grouping - always use SUM for totals
            if color_param:
                df_agg = df.groupby([x_col, color_param], as_index=False)[y_col].sum()
                # Convert to numeric and clean
                df_agg[y_col] = pd.to_numeric(df_agg[y_col], errors='coerce').fillna(0)
                fig = px.bar(df_agg, x=x_col, y=y_col, color=color_param)
            else:
                df_agg = df.groupby(x_col, as_index=False)[y_col].sum()
                df_agg[y_col] = pd.to_numeric(df_agg[y_col], errors='coerce').fillna(0)
                fig = px.bar(df_agg, x=x_col, y=y_col)
        elif chart_type == 'line':
            # Aggregate data properly by grouping X and Color columns
            if color_param:
                # Group by both X axis and Color, then sum the Y values
                df_agg = df.groupby([x_col, color_param], as_index=False)[y_col].sum()
                # Ensure data types are correct - convert to float64 explicitly for Plotly 6.x
                df_agg[y_col] = pd.to_numeric(df_agg[y_col], errors='coerce').astype('float64').fillna(0)
                # Debug output
                print(f"\nDEBUG Line Chart: {len(df_agg)} rows after aggregation")
                print(f"DEBUG Y-axis ({y_col}) range: min={df_agg[y_col].min()}, max={df_agg[y_col].max()}")
                print(f"DEBUG Sample values:\n{df_agg[[x_col, color_param, y_col]].head(10)}\n")
                
                # Sort by date if x_col looks like dates (contains /)
                if df_agg[x_col].astype(str).str.contains('/', na=False).any():
                    try:
                        # Try multiple date formats: dd/mmm, dd/mm/yyyy, etc.
                        # Add year 2000 for dates without year (01/Jan becomes 01/Jan/2000)
                        def parse_date(date_str):
                            if pd.isna(date_str):
                                return pd.NaT
                            date_str = str(date_str).strip()
                            # Try with year first
                            for fmt in ['%d/%m/%Y', '%d/%b/%Y', '%d-%m-%Y', '%d-%b-%Y']:
                                try:
                                    return pd.to_datetime(date_str, format=fmt)
                                except:
                                    pass
                            # Try without year (add 2000 as default year)
                            for fmt in ['%d/%b', '%d/%m']:
                                try:
                                    parsed = pd.to_datetime(date_str + '/2000', format=fmt + '/%Y')
                                    return parsed
                                except:
                                    pass
                            return pd.NaT
                        
                        df_agg['_temp_date'] = df_agg[x_col].apply(parse_date)
                        if df_agg['_temp_date'].notna().any():
                            df_agg = df_agg.sort_values(['_temp_date', color_param])
                            df_agg = df_agg.drop(columns=['_temp_date'])
                        else:
                            df_agg = df_agg.sort_values([color_param, x_col])
                    except Exception as e:
                        print(f"Date parsing error: {e}")
                        df_agg = df_agg.sort_values([color_param, x_col])
                else:
                    df_agg = df_agg.sort_values([color_param, x_col])
                
                # Final debug before creating chart
                print(f"DEBUG FINAL: {len(df_agg)} rows before px.line()")
                print(f"DEBUG FINAL Y range: min={df_agg[y_col].min()}, max={df_agg[y_col].max()}")
                print(f"DEBUG FINAL dtypes: {df_agg.dtypes}")
                    
                fig = px.line(df_agg, x=x_col, y=y_col, color=color_param, markers=True)
            else:
                # Group by X axis only, then sum the Y values
                df_agg = df.groupby(x_col, as_index=False)[y_col].sum()
                # Ensure data types are correct - convert to float64 explicitly for Plotly 6.x
                df_agg[y_col] = pd.to_numeric(df_agg[y_col], errors='coerce').astype('float64').fillna(0)
                
                # Sort by date if x_col looks like dates
                if df_agg[x_col].astype(str).str.contains('/', na=False).any():
                    try:
                        # Try multiple date formats
                        def parse_date(date_str):
                            if pd.isna(date_str):
                                return pd.NaT
                            date_str = str(date_str).strip()
                            # Try with year first
                            for fmt in ['%d/%m/%Y', '%d/%b/%Y', '%d-%m-%Y', '%d-%b-%Y']:
                                try:
                                    return pd.to_datetime(date_str, format=fmt)
                                except:
                                    pass
                            # Try without year (add 2000 as default)
                            for fmt in ['%d/%b', '%d/%m']:
                                try:
                                    parsed = pd.to_datetime(date_str + '/2000', format=fmt + '/%Y')
                                    return parsed
                                except:
                                    pass
                            return pd.NaT
                        
                        df_agg['_temp_date'] = df_agg[x_col].apply(parse_date)
                        if df_agg['_temp_date'].notna().any():
                            df_agg = df_agg.sort_values('_temp_date')
                            df_agg = df_agg.drop(columns=['_temp_date'])
                        else:
                            df_agg = df_agg.sort_values(x_col)
                    except Exception as e:
                        print(f"Date parsing error: {e}")
                        df_agg = df_agg.sort_values(x_col)
                else:
                    df_agg = df_agg.sort_values(x_col)
                    
                fig = px.line(df_agg, x=x_col, y=y_col, markers=True)
        elif chart_type == 'scatter':
            # Scatter plots show individual points, but can still aggregate
            if color_param:
                df_agg = df.groupby([x_col, color_param], as_index=False)[y_col].sum()
                df_agg[y_col] = pd.to_numeric(df_agg[y_col], errors='coerce').astype('float64').fillna(0)
                fig = px.scatter(df_agg, x=x_col, y=y_col, color=color_param)
            else:
                df_agg = df.groupby(x_col, as_index=False)[y_col].sum()
                df_agg[y_col] = pd.to_numeric(df_agg[y_col], errors='coerce').astype('float64').fillna(0)
                fig = px.scatter(df_agg, x=x_col, y=y_col)
        elif chart_type == 'pie':
            # Aggregate for pie chart
            df_pie = df.groupby(x_col)[y_col].sum().reset_index()
            df_pie[y_col] = pd.to_numeric(df_pie[y_col], errors='coerce').astype('float64').fillna(0)
            # Limit to top 10 slices
            df_pie = df_pie.nlargest(10, y_col)
            fig = px.pie(df_pie, names=x_col, values=y_col)
        elif chart_type == 'box':
            fig = px.box(df, x=x_col, y=y_col, color=color_param)
        elif chart_type == 'histogram':
            fig = px.histogram(df, x=x_col, y=y_col, color=color_param)
        else:
            return jsonify({'error': f'Unsupported chart type: {chart_type}'}), 400
        
        if fig:
            # Update layout for better appearance and performance
            fig.update_layout(
                template='plotly_white',
                title=f'{chart_type.upper()} Chart',
                xaxis_title=x_col,
                yaxis_title=y_col,
                height=500,
                showlegend=True if color_param else False,
                xaxis={
                    'tickangle': -45,
                    'automargin': True,
                    'tickmode': 'auto',
                    'nticks': 20
                },
                yaxis={
                    'automargin': True,
                    'exponentformat': 'none',
                    'separatethousands': True
                },
                margin=dict(l=60, r=40, t=80, b=120)
            )
            
            # Serialize using Plotly's JSON encoder to convert numpy arrays to lists
            import json
            from plotly.utils import PlotlyJSONEncoder
            chart_json = json.dumps(fig, cls=PlotlyJSONEncoder)
            
            # Debug: Verify the data is correct
            parsed = json.loads(chart_json)
            if 'data' in parsed and len(parsed['data']) > 0:
                first_trace = parsed['data'][0]
                if isinstance(first_trace.get('y'), list):
                    y_values = first_trace['y']
                    print(f"\nDEBUG JSON: First trace has {len(y_values)} Y values (list)")
                    print(f"DEBUG JSON Y range: min={min(y_values)}, max={max(y_values)}")
                    print(f"DEBUG JSON First 10 Y values: {y_values[:10]}\n")
                else:
                    print(f"\nDEBUG JSON: y type after encoding: {type(first_trace.get('y'))}")
                    print(f"Content: {first_trace.get('y')}\n")
            
            return jsonify({
                'success': True,
                'chart': chart_json  # Send as JSON string for frontend to parse
            })
        else:
            return jsonify({'error': 'Failed to create chart'}), 500
    
    except Exception as e:
        print(f"Chart error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

# Helper: build Plotly figure from chart config and dataframe (for exports)
def build_chart_figure(df, chart_config):
    """Create a Plotly figure from a dataframe and a chart config dict.
    Expected keys in chart_config: chart_type, x, y, color, filters (optional)
    """
    import plotly.express as px
    import warnings
    cfg = chart_config or {}
    chart_type = cfg.get('chart_type') or cfg.get('type') or 'line'
    x_col = cfg.get('x') or cfg.get('x_column')
    y_col = cfg.get('y') or cfg.get('y_column')
    color_col = cfg.get('color') or cfg.get('color_column')
    filters = cfg.get('filters', {})

    if df.empty or not x_col or not y_col:
        return None

    # Apply filters first
    if filters:
        for col, values in filters.items():
            if col in df.columns and values:
                df = df[df[col].isin(values)]
    if df.empty:
        return None

    # Calculated columns
    df = add_calculated_columns(df)

    # Clean data
    df[x_col] = df[x_col].fillna('').astype(str)
    df[y_col] = pd.to_numeric(df[y_col], errors='coerce').fillna(0)
    if color_col and color_col in df.columns:
        df[color_col] = df[color_col].fillna('')
    df = df.dropna(subset=[x_col, y_col])
    if df.empty:
        return None

    # Prefer date strings over week codes for line
    if chart_type == 'line':
        has_dates = df[x_col].str.contains('/', na=False).any()
        has_week_codes = df[x_col].str.match(r'^W\d+$', case=False, na=False).any()
        if has_dates and has_week_codes:
            df = df[df[x_col].str.contains('/', na=False)]
            if df.empty:
                return None

    # Sort if looks like date
    try:
        with warnings.catch_warnings():
            warnings.filterwarnings('ignore', category=UserWarning)
            x_sort_col = f'{x_col}_sort_temp'
            df[x_sort_col] = pd.to_datetime(df[x_col], errors='coerce', format='mixed')
        if df[x_sort_col].notna().any():
            df = df.sort_values(x_sort_col).drop(columns=[x_sort_col])
    except Exception:
        pass

    # Aggregate for most chart types
    color_param = color_col if color_col and color_col in df.columns else None
    fig = None
    if chart_type == 'bar':
        if color_param:
            df_agg = df.groupby([x_col, color_param], as_index=False)[y_col].sum()
            df_agg[y_col] = pd.to_numeric(df_agg[y_col], errors='coerce').fillna(0)
            fig = px.bar(df_agg, x=x_col, y=y_col, color=color_param)
        else:
            df_agg = df.groupby(x_col, as_index=False)[y_col].sum()
            df_agg[y_col] = pd.to_numeric(df_agg[y_col], errors='coerce').fillna(0)
            fig = px.bar(df_agg, x=x_col, y=y_col)
    elif chart_type == 'line':
        if color_param:
            df_agg = df.groupby([x_col, color_param], as_index=False)[y_col].sum()
            df_agg[y_col] = pd.to_numeric(df_agg[y_col], errors='coerce').astype('float64').fillna(0)
            fig = px.line(df_agg, x=x_col, y=y_col, color=color_param, markers=True)
        else:
            df_agg = df.groupby(x_col, as_index=False)[y_col].sum()
            df_agg[y_col] = pd.to_numeric(df_agg[y_col], errors='coerce').astype('float64').fillna(0)
            fig = px.line(df_agg, x=x_col, y=y_col, markers=True)
    elif chart_type == 'scatter':
        if color_param:
            df_agg = df.groupby([x_col, color_param], as_index=False)[y_col].sum()
            df_agg[y_col] = pd.to_numeric(df_agg[y_col], errors='coerce').astype('float64').fillna(0)
            fig = px.scatter(df_agg, x=x_col, y=y_col, color=color_param)
        else:
            df_agg = df.groupby(x_col, as_index=False)[y_col].sum()
            df_agg[y_col] = pd.to_numeric(df_agg[y_col], errors='coerce').astype('float64').fillna(0)
            fig = px.scatter(df_agg, x=x_col, y=y_col)
    elif chart_type == 'pie':
        df_pie = df.groupby(x_col)[y_col].sum().reset_index()
        df_pie[y_col] = pd.to_numeric(df_pie[y_col], errors='coerce').astype('float64').fillna(0)
        df_pie = df_pie.nlargest(10, y_col)
        fig = px.pie(df_pie, names=x_col, values=y_col)
    elif chart_type == 'box':
        fig = px.box(df, x=x_col, y=y_col, color=color_param)
    elif chart_type == 'histogram':
        fig = px.histogram(df, x=x_col, y=y_col, color=color_param)
    else:
        return None

    if fig:
        fig.update_layout(
            template='plotly_white',
            title=f'{chart_type.upper()} Chart',
            xaxis_title=x_col,
            yaxis_title=y_col,
            height=500,
            showlegend=True if color_param else False,
            xaxis={'tickangle': -45, 'automargin': True, 'tickmode': 'auto', 'nticks': 20},
            yaxis={'automargin': True, 'exponentformat': 'none', 'separatethousands': True},
            margin=dict(l=60, r=40, t=80, b=120)
        )
    return fig

@app.route('/api/favorites', methods=['GET'])
def get_favorites():
    """Get list of favorite reports"""
    data = load_favorites()
    return jsonify(data)

@app.route('/api/favorites', methods=['POST'])
def save_favorite():
    """Save a favorite report"""
    try:
        report_name = request.json.get('name')
        report_config = request.json.get('config')
        
        data = load_favorites()
        data['reports'][report_name] = report_config
        save_favorites(data['reports'], report_name)
        
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/favorites/<name>', methods=['DELETE'])
def delete_favorite(name):
    """Delete a favorite report"""
    try:
        data = load_favorites()
        if name in data['reports']:
            del data['reports'][name]
            last_loaded = data['_last_loaded'] if data['_last_loaded'] != name else None
            save_favorites(data['reports'], last_loaded)
            return jsonify({'success': True})
        else:
            return jsonify({'error': 'Report not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/export', methods=['POST'])
@login_required
def export_report():
    """Export comprehensive report to Word or Excel with data, charts, and pivot tables"""
    try:
        export_format = request.json.get('format', 'excel')
        filters = request.json.get('filters', {})
        pivot_config = request.json.get('pivot_config', None)
        chart_configs = request.json.get('chart_configs', [])
        
        # Get user filter
        user_filter = None if session.get('role') == 'admin' else session.get('name')
        
        # Load and filter data from database
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        df = add_calculated_columns(df)
        
        # Apply filters
        for col, values in filters.items():
            if col in df.columns and values:
                df_col_str = df[col].astype(str)
                values_str = [str(v) for v in values]
                df = df[df_col_str.isin(values_str)]
        
        if df.empty:
            return jsonify({'error': 'No data available to export'}), 400
        
        if export_format == 'word':
            # Create comprehensive Word document
            doc = Document()
            
            # Title
            title = doc.add_heading('📊 Data Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date and metadata
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
            date_run.font.size = Pt(11)
            
            user_para = doc.add_paragraph()
            user_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            user_run = user_para.add_run(f'Generated by: {session.get("name", "User")}')
            user_run.font.size = Pt(10)
            
            doc.add_page_break()
            
            # Dataset Overview
            doc.add_heading('1. Dataset Overview', 1)
            doc.add_paragraph(f'📋 Total Rows: {df.shape[0]:,}')
            doc.add_paragraph(f'📊 Total Columns: {df.shape[1]}')
            
            # Unique staff count
            if 'Name Surname' in df.columns:
                unique_staff = df['Name Surname'].nunique()
                doc.add_paragraph(f'👥 Unique Staff Members: {unique_staff}')
            
            # Active filters
            if filters:
                doc.add_heading('Active Filters:', 2)
                for col, values in filters.items():
                    doc.add_paragraph(f'• {col}: {len(values)} selection(s)', style='List Bullet')
            
            doc.add_page_break()
            
            # Sample data table (first 20 rows)
            doc.add_heading('2. Sample Data (First 20 Rows)', 1)
            
            # Limit columns for readability in Word
            display_cols = df.columns[:10].tolist()  # First 10 columns
            sample_df = df[display_cols].head(20)
            
            # Create table
            table = doc.add_table(rows=len(sample_df) + 1, cols=len(display_cols))
            table.style = 'Light Grid Accent 1'
            
            # Header row
            for i, col in enumerate(display_cols):
                table.rows[0].cells[i].text = str(col)
                table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
            
            # Data rows
            for row_idx, (_, row) in enumerate(sample_df.iterrows(), start=1):
                for col_idx, col in enumerate(display_cols):
                    cell_value = str(row[col]) if pd.notna(row[col]) else ''
                    table.rows[row_idx].cells[col_idx].text = cell_value
            
            # Pivot table with actual data
            if pivot_config:
                doc.add_page_break()
                doc.add_heading('3. Pivot Analysis', 1)
                
                # Configuration info
                doc.add_paragraph(f'📊 Index (Group By): {pivot_config.get("index", "N/A")}')
                doc.add_paragraph(f'📋 Columns: {pivot_config.get("columns", "None")}')
                doc.add_paragraph(f'📈 Values: {", ".join(pivot_config.get("values", []))}')
                doc.add_paragraph(f'🔢 Aggregation: {pivot_config.get("agg_func", "sum")}')
                doc.add_paragraph('')
                
                # Create actual pivot table
                try:
                    index_col = pivot_config.get('index')
                    columns_col = pivot_config.get('columns')
                    values_cols = pivot_config.get('values', [])
                    agg_func = pivot_config.get('agg_func', 'sum')
                    
                    if index_col and values_cols:
                        # Convert numeric columns
                        for col in values_cols:
                            if col in df.columns:
                                df[col] = pd.to_numeric(df[col], errors='coerce')
                        
                        pivot_df = pd.pivot_table(
                            df,
                            values=values_cols[0] if len(values_cols) == 1 else values_cols,
                            index=index_col,
                            columns=columns_col if columns_col else None,
                            aggfunc=agg_func,
                            fill_value=0
                        )
                        
                        # Reset index to make it a regular column
                        pivot_df = pivot_df.reset_index()
                        
                        # Limit to first 30 rows for readability
                        display_pivot = pivot_df.head(30)
                        
                        # Create Word table
                        pivot_table = doc.add_table(rows=len(display_pivot) + 1, cols=len(display_pivot.columns))
                        pivot_table.style = 'Light Grid Accent 1'
                        
                        # Header row
                        for i, col in enumerate(display_pivot.columns):
                            cell = pivot_table.rows[0].cells[i]
                            cell.text = str(col)
                            cell.paragraphs[0].runs[0].font.bold = True
                        
                        # Data rows
                        for row_idx, (_, row) in enumerate(display_pivot.iterrows(), start=1):
                            for col_idx, col in enumerate(display_pivot.columns):
                                value = row[col]
                                if isinstance(value, (int, float)):
                                    cell_text = f'{value:,.2f}' if not pd.isna(value) else '0'
                                else:
                                    cell_text = str(value) if pd.notna(value) else ''
                                pivot_table.rows[row_idx].cells[col_idx].text = cell_text
                        
                        if len(pivot_df) > 30:
                            doc.add_paragraph(f'\n(Showing first 30 of {len(pivot_df)} total rows. Full table in Excel export.)')
                except Exception as e:
                    doc.add_paragraph(f'Error creating pivot table: {str(e)}')
                    doc.add_paragraph('Full pivot table available in Excel export.')
            
            # Chart configurations and data
            if chart_configs:
                doc.add_page_break()
                doc.add_heading('4. Chart Data', 1)
                doc.add_paragraph(f'Total Charts Created: {len(chart_configs)}')
                doc.add_paragraph('')
                
                for i, config in enumerate(chart_configs, 1):
                    chart_type = config.get("chart_type", "Unknown").title()
                    x_col = config.get("x_column", "")
                    y_col = config.get("y_column", "")
                    color_col = config.get("color_column", "")
                    
                    doc.add_heading(f'Chart {i}: {chart_type}', 2)
                    
                    # Configuration section
                    if chart_type.lower() == 'pie':
                        doc.add_paragraph(f'📊 Labels (Categories): {x_col}')
                        doc.add_paragraph(f'📈 Values (Sizes): {y_col}')
                    else:
                        doc.add_paragraph(f'📊 X-axis: {x_col}')
                        doc.add_paragraph(f'📈 Y-axis: {y_col}')
                    
                    if color_col:
                        doc.add_paragraph(f'🎨 Color by: {color_col}')
                    doc.add_paragraph('')

                    # Try to generate and embed chart image
                    try:
                        fig = build_chart_figure(df.copy(), {
                            'chart_type': config.get('chart_type'),
                            'x': config.get('x_column'),
                            'y': config.get('y_column'),
                            'color': config.get('color_column'),
                            'filters': config.get('filters', {})
                        })
                        if fig:
                            img_bytes = fig.to_image(format='png', width=1000, height=600, scale=2)
                            import tempfile
                            tmp_img = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                            tmp_img.write(img_bytes)
                            tmp_img.flush()
                            doc.add_picture(tmp_img.name, width=Inches(6))
                            doc.add_paragraph('')
                    except Exception as e:
                        doc.add_paragraph(f'⚠️ Chart image export failed: {str(e)}')
                        doc.add_paragraph('')
                    
                    # Generate chart data table
                    try:
                        # Prepare data for chart based on type
                        if x_col in df.columns and y_col in df.columns:
                            # Select columns based on what's configured
                            cols_to_use = [x_col, y_col]
                            
                            # For bar/pie charts, we need to aggregate by category
                            if chart_type.lower() in ['bar', 'pie']:
                                chart_df = df[[x_col, y_col]].copy()
                                chart_df[y_col] = pd.to_numeric(chart_df[y_col], errors='coerce')
                                # Aggregate by X column (category)
                                chart_data = chart_df.groupby(x_col)[y_col].sum().reset_index()
                                chart_data.columns = [x_col, f'Total {y_col}']
                                chart_data = chart_data.sort_values(by=f'Total {y_col}', ascending=False).head(25)
                            
                            # For line charts, show time series data
                            elif chart_type.lower() == 'line':
                                chart_df = df[[x_col, y_col]].copy()
                                chart_df[y_col] = pd.to_numeric(chart_df[y_col], errors='coerce')
                                # Group by X column (time) and aggregate
                                chart_data = chart_df.groupby(x_col)[y_col].sum().reset_index()
                                chart_data.columns = [x_col, f'Total {y_col}']
                                # Try to sort by date if possible
                                try:
                                    chart_data[x_col] = pd.to_datetime(chart_data[x_col], errors='coerce')
                                    chart_data = chart_data.sort_values(by=x_col)
                                    chart_data[x_col] = chart_data[x_col].dt.strftime('%d/%b/%Y')
                                except:
                                    pass
                                chart_data = chart_data.head(25)
                            
                            # For scatter and other types, show raw data points
                            else:
                                chart_df = df[[x_col, y_col]].copy()
                                chart_df[y_col] = pd.to_numeric(chart_df[y_col], errors='coerce')
                                chart_data = chart_df.dropna().head(50)
                            
                            # Create Word table for chart data
                            if not chart_data.empty:
                                doc.add_paragraph('Data Summary:', style='Heading 3')
                                
                                # Limit rows for display
                                display_chart = chart_data.head(25)
                                chart_table = doc.add_table(rows=len(display_chart) + 1, cols=len(display_chart.columns))
                                chart_table.style = 'Light Grid Accent 1'
                                
                                # Header row
                                for col_idx, col in enumerate(display_chart.columns):
                                    cell = chart_table.rows[0].cells[col_idx]
                                    cell.text = str(col)
                                    cell.paragraphs[0].runs[0].font.bold = True
                                
                                # Data rows
                                for row_idx, (_, row) in enumerate(display_chart.iterrows(), start=1):
                                    for col_idx, col in enumerate(display_chart.columns):
                                        value = row[col]
                                        if isinstance(value, (int, float)):
                                            cell_text = f'{value:,.2f}' if not pd.isna(value) else '0'
                                        else:
                                            cell_text = str(value) if pd.notna(value) else ''
                                        chart_table.rows[row_idx].cells[col_idx].text = cell_text
                                
                                if len(chart_data) > 25:
                                    doc.add_paragraph(f'(Showing top 25 of {len(chart_data)} data points)')
                                
                                doc.add_paragraph('')
                            else:
                                doc.add_paragraph('⚠️ No data available for this chart.')
                        else:
                            doc.add_paragraph(f'⚠️ Missing columns: {x_col} or {y_col}')
                    except Exception as e:
                        doc.add_paragraph(f'⚠️ Error generating chart data: {str(e)}')
                    
                    doc.add_paragraph('')
                
                doc.add_paragraph('Note: Interactive chart visualizations are available in the web application. ')
                doc.add_paragraph('This report includes the underlying chart data for analysis.')
            
            # Save Word document
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            return send_file(
                doc_io,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            )
        
        else:  # Excel export with multiple sheets and chart images
            excel_io = io.BytesIO()
            with pd.ExcelWriter(excel_io, engine='xlsxwriter') as writer:
                workbook = writer.book
                
                # Sheet 1: Full filtered data
                df.to_excel(writer, sheet_name='Filtered Data', index=False)
                
                # Format the data sheet
                worksheet = writer.sheets['Filtered Data']
                for i, col in enumerate(df.columns):
                    max_len = max(df[col].astype(str).apply(len).max(), len(str(col))) + 2
                    worksheet.set_column(i, i, min(max_len, 50))
                
                # Sheet 2: Summary statistics
                summary_df = df.describe(include='all').transpose()
                summary_df.to_excel(writer, sheet_name='Summary Statistics')
                
                # Sheet 3: Pivot table (if config provided)
                if pivot_config:
                    try:
                        index_col = pivot_config.get('index')
                        columns_col = pivot_config.get('columns')
                        values_cols = pivot_config.get('values', [])
                        agg_func = pivot_config.get('agg_func', 'sum')
                        
                        if index_col and values_cols:
                            pivot_df = pd.pivot_table(
                                df,
                                values=values_cols[0] if len(values_cols) == 1 else values_cols,
                                index=index_col,
                                columns=columns_col if columns_col else None,
                                aggfunc=agg_func,
                                fill_value=0
                            )
                            pivot_df.to_excel(writer, sheet_name='Pivot Table')
                    except Exception as e:
                        print(f"Error creating pivot in Excel: {e}")
                
                # Sheet 4: Metadata
                metadata = {
                    'Export Date': [datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
                    'Generated By': [session.get('name', 'User')],
                    'Total Rows': [df.shape[0]],
                    'Total Columns': [df.shape[1]],
                    'Filters Applied': [len(filters)]
                }
                metadata_df = pd.DataFrame(metadata)
                metadata_df.to_excel(writer, sheet_name='Metadata', index=False)
                
                # Sheet 5: Chart configurations (if provided)
                if chart_configs:
                    chart_data = []
                    for i, config in enumerate(chart_configs, 1):
                        chart_data.append({
                            'Chart Number': i,
                            'Type': config.get('chart_type', ''),
                            'X-axis': config.get('x_column', ''),
                            'Y-axis': config.get('y_column', ''),
                            'Color': config.get('color_column', 'None')
                        })
                    charts_df = pd.DataFrame(chart_data)
                    charts_df.to_excel(writer, sheet_name='Chart Configs', index=False)

                    # Charts sheet with embedded PNGs
                    charts_ws = writer.book.add_worksheet('Charts')
                    row = 0
                    col = 0
                    for i, config in enumerate(chart_configs, 1):
                        try:
                            fig = build_chart_figure(df.copy(), {
                                'chart_type': config.get('chart_type'),
                                'x': config.get('x_column'),
                                'y': config.get('y_column'),
                                'color': config.get('color_column'),
                                'filters': config.get('filters', {})
                            })
                            if fig:
                                img_bytes = fig.to_image(format='png', width=1000, height=600, scale=2)
                                import tempfile
                                tmp_img = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                                tmp_img.write(img_bytes)
                                tmp_img.flush()
                                charts_ws.write(row, col, f'Chart {i}: {config.get("chart_type", "").title()}')
                                charts_ws.insert_image(row + 1, col, tmp_img.name, {'x_scale': 0.7, 'y_scale': 0.7})
                                row += 35
                        except Exception as e:
                            charts_ws.write(row, col, f'Chart {i} export failed: {str(e)}')
                            row += 2
            
            excel_io.seek(0)
            
            return send_file(
                excel_io,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=f'report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
            )
    
    except Exception as e:
        print(f"Export error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/export-pivot', methods=['POST'])
@login_required
def export_pivot():
    """Export pivot table to Excel or Word"""
    try:
        print("=== PIVOT EXPORT DEBUG ===")
        export_format = request.json.get('format', 'excel')
        filters = request.json.get('filters', {})
        pivot_config = request.json.get('pivot_config', None)
        
        print(f"Export format: {export_format}")
        print(f"Pivot config: {pivot_config}")
        
        if not pivot_config:
            return jsonify({'error': 'No pivot configuration provided'}), 400
        
        # Get user filter
        user_filter = None if session.get('role') == 'admin' else session.get('name')
        
        # Load and filter data
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        df = add_calculated_columns(df)
        
        # Apply filters
        for col, values in filters.items():
            if col in df.columns and values:
                df_col_str = df[col].astype(str)
                values_str = [str(v) for v in values]
                df = df[df_col_str.isin(values_str)]
        
        if df.empty:
            return jsonify({'error': 'No data available to export'}), 400
        
        # Create pivot table
        index_col = pivot_config.get('index')
        columns_col = pivot_config.get('columns')
        values_cols = pivot_config.get('values')
        agg_func = pivot_config.get('agg_func', 'sum')
        
        # Handle both single value and list of values
        if isinstance(values_cols, str):
            values_cols = [values_cols]
        elif not isinstance(values_cols, list):
            values_cols = [str(values_cols)] if values_cols else []
        
        print(f"Index: {index_col}, Columns: {columns_col}, Values: {values_cols}, Agg: {agg_func}")
        
        if not index_col or not values_cols:
            return jsonify({'error': 'Invalid pivot configuration - missing index or values'}), 400
        
        # Build pivot
        try:
            # Validate and convert all value columns to numeric
            valid_values = []
            for val_col in values_cols:
                if val_col in df.columns:
                    df[val_col] = pd.to_numeric(df[val_col], errors='coerce')
                    valid_values.append(val_col)
                    print(f"Converted {val_col} to numeric, non-null count: {df[val_col].notna().sum()}")
                else:
                    print(f"Warning: Value column '{val_col}' not found in data")
            
            if not valid_values:
                return jsonify({'error': 'None of the selected value columns were found in the data'}), 400
            
            print(f"Using value columns: {valid_values}")
            print(f"Available columns: {df.columns.tolist()[:10]}")
            
            # Use all valid value columns (or just first if only one)
            values_to_use = valid_values[0] if len(valid_values) == 1 else valid_values
            
            pivot_params = {
                'values': values_to_use,
                'index': index_col,
                'aggfunc': agg_func,
                'fill_value': 0
            }
            
            if columns_col and columns_col in df.columns:
                pivot_params['columns'] = columns_col
                print(f"Added columns parameter: {columns_col}")
            
            pivot_df = pd.pivot_table(df, **pivot_params)
            pivot_df = pivot_df.reset_index()
            
            # Handle multi-level columns
            if any(isinstance(col, tuple) for col in pivot_df.columns):
                pivot_df.columns = [' - '.join(map(str, col)).strip(' -') if isinstance(col, tuple) else str(col) for col in pivot_df.columns]
            
        except Exception as e:
            return jsonify({'error': f'Error creating pivot: {str(e)}'}), 500
        
        if export_format == 'word':
            # Create Word document with pivot table
            doc = Document()
            
            # Title
            title = doc.add_heading('📊 Pivot Table Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
            date_run.font.size = Pt(11)
            
            doc.add_paragraph()
            
            # Configuration
            doc.add_heading('Pivot Configuration', 1)
            doc.add_paragraph(f'📊 Rows (Index): {index_col}')
            if columns_col:
                doc.add_paragraph(f'📋 Columns: {columns_col}')
            doc.add_paragraph(f'📈 Values: {", ".join(valid_values)}')
            doc.add_paragraph(f'🔢 Aggregation: {agg_func.upper()}')
            
            doc.add_paragraph()
            
            # Pivot table data
            doc.add_heading('Pivot Table', 1)
            
            # Limit to reasonable size for Word
            display_pivot = pivot_df.head(100)
            
            # Create table
            pivot_table = doc.add_table(rows=len(display_pivot) + 1, cols=len(display_pivot.columns))
            pivot_table.style = 'Light Grid Accent 1'
            
            # Header row
            for i, col in enumerate(display_pivot.columns):
                cell = pivot_table.rows[0].cells[i]
                cell.text = str(col)
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Data rows
            for row_idx, (_, row) in enumerate(display_pivot.iterrows(), start=1):
                for col_idx, col in enumerate(display_pivot.columns):
                    value = row[col]
                    if isinstance(value, (int, float)):
                        cell_text = f'{value:,.2f}' if not pd.isna(value) else '0'
                    else:
                        cell_text = str(value) if pd.notna(value) else ''
                    pivot_table.rows[row_idx].cells[col_idx].text = cell_text
            
            if len(pivot_df) > 100:
                doc.add_paragraph(f'\n(Showing first 100 of {len(pivot_df)} total rows)')
            
            # Save Word document
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            return send_file(
                doc_io,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'pivot_table_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            )
        
        else:  # Excel export
            excel_io = io.BytesIO()
            with pd.ExcelWriter(excel_io, engine='xlsxwriter') as writer:
                workbook = writer.book
                
                # Sheet 1: Pivot table
                pivot_df.to_excel(writer, sheet_name='Pivot Table', index=False)
                
                # Format the pivot sheet
                worksheet = writer.sheets['Pivot Table']
                for i, col in enumerate(pivot_df.columns):
                    max_len = max(pivot_df[col].astype(str).apply(len).max(), len(str(col))) + 2
                    worksheet.set_column(i, i, min(max_len, 50))
                
                # Sheet 2: Filtered data (source)
                df.to_excel(writer, sheet_name='Source Data', index=False)
                
                # Sheet 3: Configuration
                config_df = pd.DataFrame({
                    'Setting': ['Index/Rows', 'Columns', 'Values', 'Aggregation'],
                    'Value': [index_col, columns_col or 'None', ', '.join(valid_values), agg_func.upper()]
                })
                config_df.to_excel(writer, sheet_name='Configuration', index=False)
            
            excel_io.seek(0)
            
            return send_file(
                excel_io,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=f'pivot_table_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
            )
    
    except Exception as e:
        print(f"Pivot export error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/export-charts', methods=['POST'])
@login_required
def export_charts():
    """Export chart data to Excel or Word"""
    try:
        export_format = request.json.get('format', 'excel')
        filters = request.json.get('filters', {})
        chart_configs = request.json.get('chart_configs', [])
        
        if not chart_configs:
            return jsonify({'error': 'No chart configurations provided'}), 400
        
        # Get user filter
        user_filter = None if session.get('role') == 'admin' else session.get('name')
        
        # Load and filter data
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        df = add_calculated_columns(df)
        
        # Apply filters
        for col, values in filters.items():
            if col in df.columns and values:
                df_col_str = df[col].astype(str)
                values_str = [str(v) for v in values]
                df = df[df_col_str.isin(values_str)]
        
        if df.empty:
            return jsonify({'error': 'No data available to export'}), 400
        
        if export_format == 'word':
            # Create Word document with chart images
            doc = Document()
            
            # Title
            title = doc.add_heading('📈 Charts Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
            date_run.font.size = Pt(11)
            
            doc.add_paragraph()
            doc.add_heading(f'Total Charts: {len(chart_configs)}', 1)
            doc.add_paragraph()
            
            # Process each chart
            for i, config in enumerate(chart_configs, 1):
                chart_type = config.get("chart_type", "Unknown").title()
                x_col = config.get("x_column", "")
                y_col = config.get("y_column", "")
                color_col = config.get("color_column", "")
                
                doc.add_heading(f'Chart {i}: {chart_type}', 2)
                
                # Configuration
                if chart_type.lower() == 'pie':
                    doc.add_paragraph(f'📊 Categories: {x_col}')
                    doc.add_paragraph(f'📈 Values: {y_col}')
                else:
                    doc.add_paragraph(f'📊 X-axis: {x_col}')
                    doc.add_paragraph(f'📈 Y-axis: {y_col}')
                
                if color_col:
                    doc.add_paragraph(f'🎨 Color by: {color_col}')
                
                doc.add_paragraph()
                
                # Generate chart image or data table
                image_added = False
                chart_data = None
                
                try:
                    if x_col in df.columns and y_col in df.columns:
                        # Prepare chart data
                        chart_df = df[[x_col, y_col]].copy()
                        chart_df[y_col] = pd.to_numeric(chart_df[y_col], errors='coerce')
                        chart_df = chart_df.dropna()
                        
                        # Aggregate based on chart type
                        if chart_type.lower() in ['bar', 'pie']:
                            chart_data = chart_df.groupby(x_col)[y_col].sum().reset_index()
                            chart_data = chart_data.sort_values(by=y_col, ascending=False).head(30)
                        elif chart_type.lower() == 'line':
                            chart_data = chart_df.groupby(x_col)[y_col].sum().reset_index()
                            # Try to sort by date
                            try:
                                chart_data[x_col] = pd.to_datetime(chart_data[x_col], errors='coerce')
                                chart_data = chart_data.sort_values(by=x_col)
                                # Convert back to string for display
                                chart_data[x_col] = chart_data[x_col].dt.strftime('%d/%b/%Y')
                            except:
                                pass
                            chart_data = chart_data.head(50)
                        else:
                            chart_data = chart_df.head(50)
                        
                        # Try to create chart image
                        try:
                            # Create Plotly figure
                            fig = None
                            if chart_type.lower() == 'bar':
                                fig = px.bar(chart_data, x=x_col, y=y_col, title=f'{chart_type} Chart: {y_col} by {x_col}')
                            elif chart_type.lower() == 'line':
                                fig = px.line(chart_data, x=x_col, y=y_col, title=f'{chart_type} Chart: {y_col} over {x_col}')
                            elif chart_type.lower() == 'pie':
                                fig = px.pie(chart_data, names=x_col, values=y_col, title=f'{chart_type} Chart: {y_col} by {x_col}')
                            elif chart_type.lower() == 'scatter':
                                fig = px.scatter(chart_data, x=x_col, y=y_col, title=f'{chart_type} Chart: {y_col} vs {x_col}')
                            
                            if fig:
                                # Update layout for better export
                                fig.update_layout(
                                    width=800,
                                    height=500,
                                    template='plotly_white',
                                    showlegend=True
                                )
                                
                                # Export to image using kaleido
                                img_bytes = fig.to_image(format='png', width=800, height=500, scale=2)
                                img_stream = io.BytesIO(img_bytes)
                                doc.add_picture(img_stream, width=Inches(6))
                                image_added = True
                                doc.add_paragraph()
                        except Exception as img_error:
                            print(f"Chart image generation failed: {str(img_error)}")
                            # Will fall back to table below
                        
                        # If image failed, show data table as fallback
                        if not image_added and chart_data is not None and not chart_data.empty:
                            doc.add_paragraph('Chart Data Table:', style='Heading 3')
                            
                            # Create table
                            table = doc.add_table(rows=min(len(chart_data) + 1, 31), cols=2)
                            table.style = 'Light Grid Accent 1'
                            
                            # Header
                            table.rows[0].cells[0].text = str(x_col)
                            table.rows[0].cells[1].text = str(y_col)
                            for cell in table.rows[0].cells:
                                cell.paragraphs[0].runs[0].font.bold = True
                            
                            # Data rows (max 30)
                            for idx, (_, row) in enumerate(chart_data.head(30).iterrows(), start=1):
                                table.rows[idx].cells[0].text = str(row[x_col]) if pd.notna(row[x_col]) else ''
                                value = row[y_col]
                                table.rows[idx].cells[1].text = f'{value:,.2f}' if pd.notna(value) else '0'
                            
                            doc.add_paragraph()
                    else:
                        doc.add_paragraph(f'⚠️ Missing columns: {x_col} or {y_col}')
                
                except Exception as e:
                    doc.add_paragraph(f'⚠️ Error generating chart: {str(e)}')
                
                doc.add_paragraph()
            
            # Save Word document
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            return send_file(
                doc_io,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'charts_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            )
        
        else:  # Excel export
            excel_io = io.BytesIO()
            with pd.ExcelWriter(excel_io, engine='xlsxwriter') as writer:
                workbook = writer.book
                
                # Sheet 1: Chart configurations
                chart_config_data = []
                for i, config in enumerate(chart_configs, 1):
                    chart_config_data.append({
                        'Chart #': i,
                        'Type': config.get('chart_type', ''),
                        'X-axis': config.get('x_column', ''),
                        'Y-axis': config.get('y_column', ''),
                        'Color': config.get('color_column', 'None')
                    })
                config_df = pd.DataFrame(chart_config_data)
                config_df.to_excel(writer, sheet_name='Chart Configurations', index=False)
                
                # Sheets 2+: Data for each chart
                for i, config in enumerate(chart_configs, 1):
                    try:
                        x_col = config.get("x_column", "")
                        y_col = config.get("y_column", "")
                        chart_type = config.get("chart_type", "").lower()
                        
                        if x_col in df.columns and y_col in df.columns:
                            chart_df = df[[x_col, y_col]].copy()
                            chart_df[y_col] = pd.to_numeric(chart_df[y_col], errors='coerce')
                            
                            # Aggregate based on type
                            if chart_type in ['bar', 'pie']:
                                chart_data = chart_df.groupby(x_col)[y_col].sum().reset_index()
                                chart_data.columns = [x_col, f'Total {y_col}']
                                chart_data = chart_data.sort_values(by=f'Total {y_col}', ascending=False)
                            elif chart_type == 'line':
                                chart_data = chart_df.groupby(x_col)[y_col].sum().reset_index()
                                chart_data.columns = [x_col, f'Total {y_col}']
                            else:
                                chart_data = chart_df.dropna().head(100)
                            
                            # Save to sheet
                            sheet_name = f'Chart {i} Data'[:31]  # Excel sheet name limit
                            chart_data.to_excel(writer, sheet_name=sheet_name, index=False)
                    except:
                        continue
                
                # Last sheet: Filtered source data
                df.to_excel(writer, sheet_name='Source Data', index=False)
            
            excel_io.seek(0)
            
            return send_file(
                excel_io,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=f'charts_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
            )
    
    except Exception as e:
        print(f"Charts export error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/save-filter', methods=['POST'])
@login_required
def save_filter():
    """Save a filter configuration for the current user"""
    try:
        data = request.json
        filter_name = data.get('filter_name', '').strip()
        filter_type = data.get('filter_type', '').strip()  # 'database', 'pivot', or 'graph'
        filter_config = data.get('filter_config', {})
        
        if not filter_name:
            return jsonify({'error': 'Filter name is required'}), 400
        
        if filter_type not in ['database', 'pivot', 'graph']:
            return jsonify({'error': 'Invalid filter type. Must be database, pivot, or graph'}), 400
        
        # Get current user
        user = User.query.filter_by(username=session.get('user')).first()
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        # Check if filter with same name and type already exists for this user
        existing_filter = SavedFilter.query.filter_by(
            user_id=user.id,
            filter_name=filter_name,
            filter_type=filter_type
        ).first()
        
        if existing_filter:
            # Update existing filter
            existing_filter.filter_config = json.dumps(filter_config)
            existing_filter.updated_at = datetime.utcnow()
            db.session.commit()
            return jsonify({
                'success': True,
                'message': 'Filter updated successfully',
                'filter_id': existing_filter.id
            })
        else:
            # Create new filter
            new_filter = SavedFilter(
                user_id=user.id,
                filter_name=filter_name,
                filter_type=filter_type,
                filter_config=json.dumps(filter_config)
            )
            db.session.add(new_filter)
            db.session.commit()
            return jsonify({
                'success': True,
                'message': 'Filter saved successfully',
                'filter_id': new_filter.id
            })
    
    except Exception as e:
        db.session.rollback()
        print(f"Error saving filter: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/load-filters', methods=['GET'])
@login_required
def load_filters():
    """Load all saved filters for the current user"""
    try:
        filter_type = request.args.get('filter_type')  # Optional: filter by type
        
        # Get current user
        user = User.query.filter_by(username=session.get('user')).first()
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        # Query filters
        query = SavedFilter.query.filter_by(user_id=user.id)
        if filter_type and filter_type in ['database', 'pivot', 'graph']:
            query = query.filter_by(filter_type=filter_type)
        
        filters = query.order_by(SavedFilter.updated_at.desc()).all()
        
        # Format response
        filters_list = []
        for f in filters:
            filters_list.append({
                'id': f.id,
                'filter_name': f.filter_name,
                'filter_type': f.filter_type,
                'filter_config': json.loads(f.filter_config),
                'created_at': f.created_at.isoformat(),
                'updated_at': f.updated_at.isoformat()
            })
        
        return jsonify({
            'success': True,
            'filters': filters_list
        })
    
    except Exception as e:
        print(f"Error loading filters: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/delete-filter/<int:filter_id>', methods=['DELETE'])
@login_required
def delete_filter(filter_id):
    """Delete a saved filter"""
    try:
        # Get current user
        user = User.query.filter_by(username=session.get('user')).first()
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        # Get filter and verify ownership
        saved_filter = SavedFilter.query.filter_by(id=filter_id, user_id=user.id).first()
        if not saved_filter:
            return jsonify({'error': 'Filter not found or you do not have permission to delete it'}), 404
        
        db.session.delete(saved_filter)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Filter deleted successfully'
        })
    
    except Exception as e:
        db.session.rollback()
        print(f"Error deleting filter: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    with app.app_context():
        init_db()
    app.run(debug=True, host='0.0.0.0', port=5000)




